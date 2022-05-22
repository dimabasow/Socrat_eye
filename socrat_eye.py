
def graph (GrName, get_x, get_y, lablex="t, s", labley="", x1=None, y1=None, x2=None, y2=None, mult_x=None, mult_y=None, shift_x=None, shift_y=None, stpx=None, stpy=None):
    import matplotlib.pyplot as plt
    import matplotlib.ticker as ticker
    import pylab
    import os
    import pandas
    import numpy
    
    
    if isinstance(get_y[0],(numpy.ndarray,pandas.Series,list,tuple)):
        y=[pandas.Series(i) for i in get_y]
    else:
        y=[pandas.Series(get_y)]
    if isinstance(get_x[0],(numpy.ndarray,pandas.Series,list,tuple)):
        x=[pandas.Series(i) for i in get_x]
    else:
        x=[pandas.Series(get_x) for i in range(len(y))]
    
    colors = ('#000000', '#FF0000', '#0000FF', '#008000', '#FF00FF', '#8B4513',
         '#808080', '#C0C0C0', '#800080', '#800000', '#FFFF00', '#00FF00',
         '#000080')
    markers=('D', '^', 'o', 's', '*', 'p', 'h', '+', 'x', 'd', 'v')
    
    if (labley=='T, C') :
        labley='T, \N{DEGREE SIGN}C'
    if (labley=='F, m^2') :
        labley='F, m\u00b2'
    if (labley=='G, m^3/hr'):
        labley='G, m\u00b3' + '/hr'
    if (labley=='ρ, kg/m^3'):
        labley='ρ, kg/' + 'm\u00b3'
    
    
    
    fig, ax = plt.subplots()
    path_dict={}
    
    if mult_x is not None:
        x=[i*mult_x for i in x]
    if mult_y is not None:
        y=[i*mult_y for i in y]
    if shift_x is not None:
        x=[i+shift_x for i in x]
    if shift_y is not None:
        y=[i+shift_y for i in y]
    
    if x1 is None:
        x1=min(min(i) for i in x)
    else:
        for i in range(len(x)):
            if x1>x[i][0]:
                N_min=x[i][x[i]>x1].index[0]
                x[i]=x[i][N_min:]
                y[i]=y[i][N_min:]
    if x2 is None:
        x2=max(max(i) for i in x)
    else:
        for i in range(len(x)):
            if x2<x[i][len(x[i])-1]:
                N_max=x[i][x[i]>x2].index[0]
                x[i]=x[i][:N_max]
                y[i]=y[i][:N_max]
    
    if y1 is None:
        y1=min(min(i) for i in y)
    if y2 is None:
        y2=0.1*(max(max(i) for i in y)-y1)+max(max(i) for i in y)
        
    path_dict["x_min"]=str(int(x1))
    path_dict["x_max"]=str(int(x2))
    path_dict["y_min"]=str(int(y1))
    path_dict["y_max"]=str(int(y2))

    N_points_aray=[]
    for i in range(len(x)):
        N_points = len(x[i])//50
        if N_points==0:
            N_points=1
        N_points_aray.append(N_points)
    for i in range(len(y)):
        ax.plot(x[i], y[i], colors[i], linewidth = 2, label = str(i+1), marker = markers[i], markevery = N_points_aray[i])
    
    pylab.xlim (xmin = x1, xmax=x2)
    pylab.ylim (ymin = y1, ymax=y2)
    #Настраиваем оси
    #Устанавливаем интервал делений x:
    if stpx!=None:
        ax.xaxis.set_major_locator(ticker.MultipleLocator(stpx))
        ax.xaxis.set_minor_locator(ticker.MultipleLocator(stpx/5))
    #Устанавливаем интервал делений y:
    if stpy!=None:
        ax.yaxis.set_major_locator(ticker.MultipleLocator(stpy))
        ax.yaxis.set_minor_locator(ticker.MultipleLocator(stpy/5))

    # Отрисовываем сеточные линии
    ax.grid(which = 'major', color = 'k',  linestyle = '-')
    ax.minorticks_on()

    ax.grid(which = 'minor', color ='gray', linestyle = ':')

    ax.tick_params( axis ='both', length = 10, pad = 10, labelsize = 18, labelcolor = 'black' )

    # Отрисовываем подписи
    ax.set_ylabel(labley, fontsize = 24)
    ax.set_xlabel(lablex, fontsize = 24)
    
    if len(get_y)>1:
        pylab.legend(loc='best', numpoints=1, fontsize=20).get_frame().set_edgecolor('w')
    
    fig.set_figwidth(12)
    fig.set_figheight(8)
    
    # сохранение .png
    fmt = 'png'
    path_dict[fmt]=os.path.abspath(os.curdir)+"\\"+GrName + "." +fmt
    fig.savefig('{}.{}'.format(GrName, fmt), format=fmt, bbox_inches='tight')
    # сохранение .pdf
#    fmt = 'pdf'
#    pwd = os.getcwd()
#    iPath = '{}'.format(fmt)
#    if not os.path.exists(iPath):
#        os.mkdir(iPath)
#    os.chdir(iPath)
#    path_dict[fmt]=os.path.abspath(os.curdir) + "\\" + GrName + "." +fmt
    #fig.savefig('{}.{}'.format(name, fmt), fmt='png')
#    fig.savefig('{}.{}'.format(GrName, fmt), format='pdf', bbox_inches='tight')
#    os.chdir(pwd)
    plt.close()
    
    return path_dict
    
def critic (calc,clusters,results_final):
    import os
    import docx
    import pandas

    Tw_CROD_names=["Tw_CROD_01", "Tw_CROD_02", "Tw_CROD_03", "Tw_CROD_04", "Tw_CROD_05", "Tw_CROD_06", "Tw_CROD_07", "Tw_CROD_08", "Tw_CROD_09", "Tw_CROD_10"]
    TUO2_MIN_names=["TUO2_MIN_01", "TUO2_MIN_02", "TUO2_MIN_03", "TUO2_MIN_04", "TUO2_MIN_05", "TUO2_MIN_06", "TUO2_MIN_07", "TUO2_MIN_08", "TUO2_MIN_09", "TUO2_MIN_10"]
    Tw_BAFFLE_names=["Tw_BAFFLE_01", "Tw_BAFFLE_02", "Tw_BAFFLE_03", "Tw_BAFFLE_04", "Tw_BAFFLE_05", "Tw_BAFFLE_06", "Tw_BAFFLE_07", "Tw_BAFFLE_08", "Tw_BAFFLE_09", "Tw_BAFFLE_10"]
    RO_MAX_CORE_BYPASS_names=["RO_MAX_CORE_BYPASS_01", "RO_MAX_CORE_BYPASS_02", "RO_MAX_CORE_BYPASS_03", "RO_MAX_CORE_BYPASS_04", "RO_MAX_CORE_BYPASS_05", "RO_MAX_CORE_BYPASS_06", "RO_MAX_CORE_BYPASS_07", "RO_MAX_CORE_BYPASS_08", "RO_MAX_CORE_BYPASS_09", "RO_MAX_CORE_BYPASS_10"]
    RO_MAX_01_names=["RO_MAX_01", "RO_MAX_02", "RO_MAX_03", "RO_MAX_04", "RO_MAX_05", "RO_MAX_06", "RO_MAX_07", "RO_MAX_08", "RO_MAX_09", "RO_MAX_10"]
    TH2O_CORE_BYPASS_names=["TH2O_CORE_BYPASS_01", "TH2O_CORE_BYPASS_02", "TH2O_CORE_BYPASS_03", "TH2O_CORE_BYPASS_04", "TH2O_CORE_BYPASS_05", "TH2O_CORE_BYPASS_06", "TH2O_CORE_BYPASS_07", "TH2O_CORE_BYPASS_08", "TH2O_CORE_BYPASS_09", "TH2O_CORE_BYPASS_10"]
    TH2O_MIN_names=["TH2O_MIN_01", "TH2O_MIN_02", "TH2O_MIN_03", "TH2O_MIN_04", "TH2O_MIN_05", "TH2O_MIN_06", "TH2O_MIN_07", "TH2O_MIN_08", "TH2O_MIN_09", "TH2O_MIN_10"]

    calc=calc.make_copy()
    stages={}

    calc.last_zero_parameters_transform(Tw_CROD_names)
    for i in Tw_CROD_names:
        calc.cut_bottom_transform(i,0)
    os.chdir(results_final)

    if min(calc.get_values("Tw_LGRID5"))<=0:
        stages["stage_1"]={}
        stages["stage_1"]["calc"]=calc.make_copy()
        stages["stage_1"]["calc"].cut_top_transform("Tw_LGRID5","min")
        stages["stage_1"]["row"]=stages["stage_1"]["calc"].get_row("RO_MAX_CORE_BYPASS_I9","max")
        if stages["stage_1"]["row"] is not None:
            stages["stage_1"]["calc"].graph("Stage_1", "Time", "RO_MAX_CORE_BYPASS_I9", lablex="t, s", labley="ρ, kg/m^3")
            stages["stage_1"]["graph"]=stages["stage_1"]["calc"].graph_list[0]["png"]
        else:
            del(stages["stage_1"])
        time_p1spn=calc.get_values("#")
        if time_p1spn is not None:
            stages["stage_2"]={}
            stages["stage_all"]={}
            stages["stage_2"]["calc"]=calc.make_copy()
            stages["stage_all"]["calc"]=calc.make_copy()
            stages["stage_2"]["calc"].cut_bottom_transform("Tw_LGRID5","min")
            stages["stage_2"]["calc"].cut_top_transform("#","min")
            stages["stage_2"]["row"]=stages["stage_2"]["calc"].get_row("RO_MAX_II","max")
            stages["stage_2"]["calc"].graph("Stage_2", "Time", "RO_MAX_II", lablex="t, s", labley="ρ, kg/m^3")
            stages["stage_2"]["graph"]=stages["stage_2"]["calc"].graph_list[0]["png"]
            
            time_p1runout=calc.get_values("time")
            if time_p1runout is not None:
                stages["stage_3"]={}
                stages["stage_3"]["calc"]=calc.make_copy()
                stages["stage_3"]["calc"].cut_bottom_transform("#","min")
                stages["stage_3"]["calc"].cut_top_transform("time","min")
                stages["stage_3"]["row"]=stages["stage_3"]["calc"].get_row("RO_MAX_III","max")
                stages["stage_3"]["calc"].graph("Stage_3", "Time", "RO_MAX_III", lablex="t, s", labley="ρ, kg/m^3")
                stages["stage_3"]["graph"]=stages["stage_3"]["calc"].graph_list[0]["png"]
                stages["stage_all"]["calc"].cut_top_transform("time","min")
        
    else:
        stages["stage_4"]={}
        stages["stage_4"]["calc"]=calc.make_copy()
        stages["stage_4"]["row"]=stages["stage_4"]["calc"].get_row("RO_MAX_CORE_BYPASS_I9","max")
        stages["stage_4"]["calc"].graph("stage_4", "Time", "RO_MAX_CORE_BYPASS_I9", lablex="t, s", labley="ρ, kg/m^3")
        stages["stage_4"]["graph"]=stages["stage_4"]["calc"].graph_list[0]["png"]
        
        

    for cluster in clusters:
        with pandas.ExcelWriter(cluster.split(".")[0] + ".xlsx") as writer:
            if "stage_all" in stages:
                stages["stage_all"]["calc"].get_values(calc.dia_clusters_dict[cluster]).to_excel(writer, sheet_name="Этапы 1-3", index=False)
            if "stage_1" in stages:
                stages["stage_1"]["calc"].get_values(calc.dia_clusters_dict[cluster]).to_excel(writer, sheet_name="Этап 1", index=False)
            if "stage_2" in stages:
                stages["stage_2"]["calc"].get_values(calc.dia_clusters_dict[cluster]).to_excel(writer, sheet_name="Этап 2", index=False)
            if "stage_3" in stages:
                stages["stage_3"]["calc"].get_values(calc.dia_clusters_dict[cluster]).to_excel(writer, sheet_name="Этап 3", index=False)
            if "stage_4" in stages:
                stages["stage_4"]["calc"].get_values(calc.dia_clusters_dict[cluster]).to_excel(writer, sheet_name="Этап 4", index=False)

    doc=docx.Document()
    doc.add_paragraph("Таблица 1")
    table1=doc.add_table(rows=1, cols=3)
    table1.style="Table Grid"
    table1.cell(0,0).text="Время, с"
    table1.cell(0,1).text="Плотность т/н, кг/м^3"
    table1.cell(0,2).text="Этап ТА"

    if "stage_1" in stages:
        table1.add_row()
        table1.cell(len(table1.rows)-1,0).text="{} - {}".format(str(round(min(stages["stage_1"]["calc"].get_values("Time")))),str(round(max(stages["stage_1"]["calc"].get_values("Time")))))
        table1.cell(len(table1.rows)-1,1).text="{} - {}".format(str(round(min(stages["stage_1"]["calc"].get_values("RO_MAX_CORE_BYPASS_I9")),2)),str(round(max(stages["stage_1"]["calc"].get_values("RO_MAX_CORE_BYPASS_I9")),2)))
        table1.cell(len(table1.rows)-1,2).text="1 этап"
    if "stage_2" in stages:
        table1.add_row()
        table1.cell(len(table1.rows)-1,0).text="{} - {}".format(str(round(min(stages["stage_2"]["calc"].get_values("Time")))),str(round(max(stages["stage_2"]["calc"].get_values("Time")))))
        table1.cell(len(table1.rows)-1,1).text="{} - {}".format(str(round(min(stages["stage_2"]["calc"].get_values("RO_MAX_II")),2)),str(round(max(stages["stage_2"]["calc"].get_values("RO_MAX_II")),2)))
        table1.cell(len(table1.rows)-1,2).text="2 этап"
    if "stage_3" in stages:
        table1.add_row()
        table1.cell(len(table1.rows)-1,0).text="{} - {}".format(str(round(min(stages["stage_3"]["calc"].get_values("Time")))),str(round(max(stages["stage_3"]["calc"].get_values("Time")))))
        table1.cell(len(table1.rows)-1,1).text="{} - {}".format(str(round(min(stages["stage_3"]["calc"].get_values("RO_MAX_III")),2)),str(round(max(stages["stage_3"]["calc"].get_values("RO_MAX_III")),2)))
        table1.cell(len(table1.rows)-1,2).text="3 этап"    
    if "stage_4" in stages:
        table1.add_row()
        table1.cell(len(table1.rows)-1,0).text="{} - {}".format(str(round(min(stages["stage_4"]["calc"].get_values("Time")))),str(round(max(stages["stage_4"]["calc"].get_values("Time")))))
        table1.cell(len(table1.rows)-1,1).text="{} - {}".format(str(round(min(stages["stage_4"]["calc"].get_values("RO_MAX_CORE_BYPASS_I9")),2)),str(round(max(stages["stage_4"]["calc"].get_values("RO_MAX_CORE_BYPASS_I9")),2)))
        table1.cell(len(table1.rows)-1,2).text="4 этап"

    
    if "stage_1" in stages or "stage_4" in stages:
        row_for_table=stages["stage_1"]["row"] if "stage_1" in stages else stages["stage_4"]["row"]
        doc.add_paragraph("Таблица 2, время равно " + str(row_for_table[0]) + " с")
    
        table2=doc.add_table(rows=7, cols=11)
        table2.style="Table Grid"
        table2.cell(0,0).text="Значение параметра"
        table2.cell(1,0).text="1. Температура топлива (минимальная)"
        table2.cell(2,0).text="2. Температура выгородки"
        table2.cell(3,0).text="3. Плотность теплоносителя с учётом направляющего канала (максимальная)"
        table2.cell(4,0).text="4. Плотность теплоносителя без учёта направляющего канала (максимальная)"
        table2.cell(5,0).text="5. Температура теплоносителя с учётом направляющего канала (минимальная)"
        table2.cell(6,0).text="6. Температура теплоносителя без учёта направляющего канала (минимальная)"
    
        for k in range(10):
            table2.cell(0,k+1).text=str(k+2)
            table2.cell(1,k+1).text=str(round(float(row_for_table[TUO2_MIN_names[k]])))
            table2.cell(2,k+1).text=str(round(float(row_for_table[Tw_BAFFLE_names[k]])))
            table2.cell(3,k+1).text=str(round(float(row_for_table[RO_MAX_CORE_BYPASS_names[k]]),1))
            table2.cell(4,k+1).text=str(round(float(row_for_table[RO_MAX_01_names[k]]),1))
            table2.cell(5,k+1).text=str(round(float(row_for_table[TH2O_CORE_BYPASS_names[k]]),1))
            table2.cell(6,k+1).text=str(round(float(row_for_table[TH2O_MIN_names[k]]),1))

    doc.add_paragraph("Таблица 3")
    table3=doc.add_table(rows=1, cols=8)
    table3.style="Table Grid"
    table3.cell(0,0).text="Этапы"
    table3.cell(0,1).text="Плотность т/н"
    table3.cell(0,2).text="Температура т/н"
    table3.cell(0,3).text="Время"
    table3.cell(0,4).text="Доля UO2"
    table3.cell(0,5).text="Доля ZrO2"
    table3.cell(0,6).text="Доля Zr"
    table3.cell(0,7).text="Доля стали"

    if "stage_1" in stages:
        table3.add_row()
        table3.cell(len(table3.rows)-1,0).text="Этап 1"
        table3.cell(len(table3.rows)-1,1).text=str(round(float(stages["stage_1"]["row"]["RO_MAX_CORE_BYPASS_I9"]),1))
        table3.cell(len(table3.rows)-1,3).text=str(round(float(stages["stage_1"]["row"]["Time"])))
        table3.cell(len(table3.rows)-1,4).text=str(round(float(stages["stage_1"]["row"]["UO2_NKR_FRACTION"]),2))
        table3.cell(len(table3.rows)-1,5).text=str(round(float(stages["stage_1"]["row"]["ZrO2_NKR_FRACTION"]),2))
        table3.cell(len(table3.rows)-1,6).text=str(round(float(stages["stage_1"]["row"]["Zr_NKR_FRACTION"]),2))
        table3.cell(len(table3.rows)-1,7).text=str(round(float(stages["stage_1"]["row"]["SS_NKR_FRACTION"]),2))
    if "stage_2" in stages:
        table3.add_row()
        table3.cell(len(table3.rows)-1,0).text="Этап 2"
        table3.cell(len(table3.rows)-1,1).text=str(round(float(stages["stage_2"]["row"]["RO_MAX_II"]),1))
        table3.cell(len(table3.rows)-1,2).text=str(round(float(stages["stage_2"]["row"]["TH2O_MIN_II"]),1))
        table3.cell(len(table3.rows)-1,3).text=str(round(float(stages["stage_2"]["row"]["Time"])))
        table3.cell(len(table3.rows)-1,4).text=str(round(float(stages["stage_2"]["row"]["UO2_NKR_FRACTION"]),2))
        table3.cell(len(table3.rows)-1,5).text=str(round(float(stages["stage_2"]["row"]["ZrO2_NKR_FRACTION"]),2))
        table3.cell(len(table3.rows)-1,6).text=str(round(float(stages["stage_2"]["row"]["Zr_NKR_FRACTION"]),2))
        table3.cell(len(table3.rows)-1,7).text=str(round(float(stages["stage_2"]["row"]["SS_NKR_FRACTION"]),2))
    if "stage_3" in stages:
        table3.add_row()
        table3.cell(len(table3.rows)-1,0).text="Этап 3"
        table3.cell(len(table3.rows)-1,1).text=str(round(float(stages["stage_3"]["row"]["RO_MAX_III"]),1))
        table3.cell(len(table3.rows)-1,2).text=str(round(float(stages["stage_3"]["row"]["TH2O_MIN_III"]),1))
        table3.cell(len(table3.rows)-1,3).text=str(round(float(stages["stage_3"]["row"]["Time"])))
        table3.cell(len(table3.rows)-1,4).text=str(round(float(stages["stage_3"]["row"]["UO2_NKR_FRACTION"]),2))
        table3.cell(len(table3.rows)-1,5).text=str(round(float(stages["stage_3"]["row"]["ZrO2_NKR_FRACTION"]),2))
        table3.cell(len(table3.rows)-1,6).text=str(round(float(stages["stage_3"]["row"]["Zr_NKR_FRACTION"]),2))
        table3.cell(len(table3.rows)-1,7).text=str(round(float(stages["stage_3"]["row"]["SS_NKR_FRACTION"]),2))
    if "stage_4" in stages:
        table3.add_row()
        table3.cell(len(table3.rows)-1,0).text="Этап 4"
        table3.cell(len(table3.rows)-1,1).text=str(round(float(stages["stage_4"]["row"]["RO_MAX_CORE_BYPASS_I9"]),1))
        table3.cell(len(table3.rows)-1,3).text=str(round(float(stages["stage_4"]["row"]["Time"])))
        table3.cell(len(table3.rows)-1,4).text=str(round(float(stages["stage_4"]["row"]["UO2_NKR_FRACTION"]),2))
        table3.cell(len(table3.rows)-1,5).text=str(round(float(stages["stage_4"]["row"]["ZrO2_NKR_FRACTION"]),2))
        table3.cell(len(table3.rows)-1,6).text=str(round(float(stages["stage_4"]["row"]["Zr_NKR_FRACTION"]),2))
        table3.cell(len(table3.rows)-1,7).text=str(round(float(stages["stage_4"]["row"]["SS_NKR_FRACTION"]),2))

    if "stage_1" in stages:
        doc.add_picture(stages["stage_1"]["graph"], width = docx.shared.Cm(12))
        doc.add_paragraph("Максимальная плотность в активной зоне на 1 этапе")
    if "stage_2" in stages:
        doc.add_picture(stages["stage_2"]["graph"], width = docx.shared.Cm(12))
        doc.add_paragraph("Максимальная плотность в шахте реактора на 2 этапе")
    if "stage_3" in stages:
        doc.add_picture(stages["stage_3"]["graph"], width = docx.shared.Cm(12))
        doc.add_paragraph("Максимальная плотность на днище корпуса реактора на 3 этапе")
    if "stage_4" in stages:
        doc.add_picture(stages["stage_4"]["graph"], width = docx.shared.Cm(12))
        doc.add_paragraph("Максимальная плотность в активной зоне")

    doc.save("Критика.docx")

class excel_input:
    def __init__(self,excel_read):
        import pandas
        excel_read.dropna(how="all",inplace=True)
        excel_read.reset_index(drop=True, inplace=True)
        excel_read=excel_read.T
        excel_read.columns=excel_read.iloc[0,:]
        excel_read.drop([0], inplace=True)
        excel_read.reset_index(drop=True, inplace=True)
        
        results_src="Путь к папке с результатами"
        results_post="Путь к обработанным результатам"
        results_final="Папка для сохранения результатов"
        critic_clusters="Кластеры для сохранения в Excel"
        not_zero_parameters="Параметры, которые не должны падать в ноль"
        last_zero_parameters="Параметры, которые если упали в ноль, не увеличиваются"
        corrected_parameters="Корректируемые параметры"
        corrected_parameters_mult="Множитель"
        corrected_parameters_shift="Смещение"
        time_parameters="Параметры для хронологии"
        time_parameters_value="Значение(число,ON,OFF,min,max)"
        time_parameters_description="Описание параметров для хронологии"
        key_parameters="Ключевые параметры"
        key_parameters_value="Значение(min,max,first,last)"
        key_parameters_description="Описание ключевых параметров"
        cut_bottom_parameters_name="Параметр для обрезки снизу"
        cut_bottom_parameters_value="Значение для обрезки снизу"
        cut_top_parameters_name="Параметр для обрезки сверху"
        cut_top_parameters_value="Значение для обрезки сверху"
        P_CORE_OUT_name="Датчик давления на выходе из АЗ"
        
        x_names="Ось X"
        y_names="Ось Y"
        lables="Подписи осей"
        cut_x="Обрезка x"
        cut_y="Обрезка y"
        mult_x_y="Умножение значений x и y"
        step_x_y="Шаг сетки по x и y"
        description_figure="Подпись к рисунку"
        
        
        self.results_src=None
        self.results_post=None
        self.results_final=None
        self.critic_clusters=None
        self.not_zero_parameters=None
        self.last_zero_parameters=None
        self.corrected_parameters=None
        self.time_parameters=None
        self.key_parameters=None
        self.graphs=None
        self.cut_bottom_parameters=None
        self.cut_top_parameters=None
        self.P_CORE_OUT_name=None
        
        
        if results_src in excel_read:
            self.results_src=list(excel_read[results_src].dropna(how="all"))
            self.results_src=[i+"\\" if i[-1]!="\\" else i for i in self.results_src]
        if results_post in excel_read:
            self.results_post=list(excel_read[results_post].dropna(how="all"))
            self.results_post=[i+"\\" if i[-1]!="\\" else i for i in self.results_post]
        if results_final in excel_read:
            self.results_final=list(excel_read[results_final].dropna(how="all"))
            self.results_final=[i+"\\" if i[-1]!="\\" else i for i in self.results_final]
        if critic_clusters in excel_read:
            self.critic_clusters=list(excel_read[critic_clusters].dropna(how="all"))
        if not_zero_parameters in excel_read:
            self.not_zero_parameters=list(excel_read[not_zero_parameters].dropna(how="all"))
        if last_zero_parameters in excel_read:
            self.last_zero_parameters=list(excel_read[last_zero_parameters].dropna(how="all"))
        if P_CORE_OUT_name in excel_read:
            self.P_CORE_OUT_name=list(excel_read[P_CORE_OUT_name].dropna(how="all"))[0]
        
        
        if corrected_parameters in excel_read:
            self.corrected_parameters={}
            corrected_parameters_names=list(excel_read[corrected_parameters].dropna(how="all"))
            for parameter in corrected_parameters_names:
                self.corrected_parameters[parameter]={}
            if corrected_parameters_mult in excel_read:
                for k in range(len(corrected_parameters_names)):
                    if pandas.isna(excel_read[corrected_parameters_mult][k]):
                        self.corrected_parameters[corrected_parameters_names[k]]["mult"]=1
                    else:
                        self.corrected_parameters[corrected_parameters_names[k]]["mult"]=excel_read[corrected_parameters_mult][k]
            if corrected_parameters_shift in excel_read:
                for k in range(len(corrected_parameters_names)):
                    if pandas.isna(excel_read[corrected_parameters_shift][k]):
                        self.corrected_parameters[corrected_parameters_names[k]]["shift"]=0
                    else:
                        self.corrected_parameters[corrected_parameters_names[k]]["shift"]=excel_read[corrected_parameters_shift][k]
        
        if cut_bottom_parameters_name in excel_read and not pandas.isna(excel_read[cut_bottom_parameters_name][0]):
            self.cut_bottom_parameters={}
            self.cut_bottom_parameters["name"]=excel_read[cut_bottom_parameters_name][0]
            self.cut_bottom_parameters["value"]=excel_read[cut_bottom_parameters_value][0]
        if cut_top_parameters_name in excel_read and not pandas.isna(excel_read[cut_top_parameters_name][0]):
            self.cut_top_parameters={}
            self.cut_top_parameters["name"]=excel_read[cut_top_parameters_name][0]
            self.cut_top_parameters["value"]=excel_read[cut_top_parameters_value][0]
        
        if time_parameters in excel_read:
            self.time_parameters=[]
            time_parameters_names=list(excel_read[time_parameters].dropna(how="all"))
            for k in range(len(time_parameters_names)):
                time_point={"name":time_parameters_names[k], "description":excel_read[time_parameters_description][k]}
                if excel_read[time_parameters_value][k]=="ON":
                    time_point["value"]=True
                elif excel_read[time_parameters_value][k]=="OFF":
                    time_point["value"]=False
                else:
                    time_point["value"]=excel_read[time_parameters_value][k]
                self.time_parameters.append(time_point)
        
        if key_parameters in excel_read:
            self.key_parameters=[]
            key_parameters_names=list(excel_read[key_parameters].dropna(how="all"))
            for k in range(len(key_parameters_names)):
                key_point={"name":key_parameters_names[k], 
                           "description":excel_read[key_parameters_description][k],
                           "value":excel_read[key_parameters_value][k]}
                self.key_parameters.append(key_point)
        
        
        if y_names in excel_read:
            self.graphs=[]
            for i in range(len(excel_read[y_names].columns)):
                graph={}
                
                graph["x"]=list(excel_read[x_names].iloc[:,i].dropna())
                if len(graph["x"])==1:
                    graph["x"]=graph["x"][0]
                
                graph["y"]=list(excel_read[y_names].iloc[:,i].dropna())
                
                graph["label_x"]=excel_read[lables].iloc[0,i]
                if pandas.isna(graph["label_x"]):
                    graph["label_x"]=""
                
                graph["label_y"]=excel_read[lables].iloc[1,i]
                if pandas.isna(graph["label_y"]):
                    graph["label_y"]=""
                
                graph["min_x"]=excel_read[cut_x].iloc[0,i]
                if pandas.isna(graph["min_x"]) or graph["min_x"]=="None":
                    graph["min_x"]=None
                
                graph["max_x"]=excel_read[cut_x].iloc[1,i]
                if pandas.isna(graph["max_x"]) or graph["max_x"]=="None":
                    graph["max_x"]=None
                
                graph["min_y"]=excel_read[cut_y].iloc[0,i]
                if pandas.isna(graph["min_y"]) or graph["min_y"]=="None":
                    graph["min_y"]=None
                
                graph["max_y"]=excel_read[cut_y].iloc[1,i]
                if pandas.isna(graph["max_y"]) or graph["max_y"]=="None":
                    graph["max_y"]=None
                
                graph["mult_x"]=excel_read[mult_x_y].iloc[0,i]
                if pandas.isna(graph["mult_x"]) or graph["mult_x"]=="None":
                    graph["mult_x"]=None
                
                graph["mult_y"]=excel_read[mult_x_y].iloc[1,i]
                if pandas.isna(graph["mult_y"]) or graph["mult_y"]=="None":
                    graph["mult_y"]=None
                
                graph["step_x"]=excel_read[step_x_y].iloc[0,i]
                if pandas.isna(graph["step_x"]) or graph["step_x"]=="None":
                    graph["step_x"]=None
                
                graph["step_y"]=excel_read[step_x_y].iloc[1,i]
                if pandas.isna(graph["step_y"]) or graph["step_y"]=="None":
                    graph["step_y"]=None
                
                graph["description"]=list(excel_read[description_figure].iloc[:,i].dropna())[0]
                
                self.graphs.append(graph)

class socrat_calculation:
    
    def __init__(self,path_to_folder, names_last_zero_parameters=[], names_not_ziro_parameters=[], dia_names=[], time_shift=None):
        import os
        import pandas
        initial_dir=os.getcwd()
        self.path_to_results=path_to_folder
        #метод, выполняемый при инициализации для парсинга файла #_report.lst
        def parsing_report(filename):
            f = open(filename,'r')
            f=f.read()
            f=f.split("\n")
            indexes=[]
            for i in range(len(f)):
                if "TRIP_ONE IS" in f[i] or "TRIP_TWO IS" in f[i] or "COMMAND PERFORMED" in f[i] or "COMMAND:" in f[i] or "COMMAND PROCESSED" in f[i]:
                    indexes.append(i)
            indexes.append(len(f))
            
            f_lists=[]
            for i in range(len(indexes)-1):
                f_lists.append(f[indexes[i]:indexes[i+1]])
            
            report_lists=[]
            for i in range(len(f_lists)):
                if "TRIP_ONE IS" in f_lists[i][0] or "TRIP_TWO IS" in f_lists[i][0] or "COMMAND PERFORMED" in f_lists[i][0]:
                    name_string=f_lists[i][0]
                    
                    index_end_list=[]
                    for k in range(len(f_lists[i])-1):
                        if "STEP=" in f_lists[i][k+1] and "TIME=" in f_lists[i][k+1]:
                            index_end_list.append(k)
                    index_end=index_end_list[0]
                    
                    condition_string="".join(f_lists[i][1:index_end+1]).strip()
                    time_string=f_lists[i][index_end+1]
                    report_lists.append([name_string,condition_string,time_string])
            
            list_for_pandas=[]
            for event in report_lists:
                time_for_pandas=float(event[2].split("TIME=")[-1])
                name_for_pandas=event[0].split(":")[-1].split()[0]
                
                if "TRIP_ONE IS OFF" in event[0] or "TRIP_TWO IS OFF" in event[0]:
                    state_for_pandas=False
                else:
                    state_for_pandas=True
                
                if "TRIP_ONE" in event[0]:
                    type_for_pandas="TRIP_ONE"
                elif "TRIP_TWO" in event[0]:
                    type_for_pandas="TRIP_TWO"
                elif "COMMAND" in event[0]:
                    type_for_pandas="COMMAND"
                else:
                    type_for_pandas=""
                
                if "DELAY:" in event[1]:
                    delay_for_pandas=float(event[1].split("DELAY:")[-1])
                else:
                    delay_for_pandas=0
                if "CONDITION:" in event[1]:
                    condition_for_pandas=event[1].split("CONDITION:")[-1].split("EXECUTION")[0]
                else:
                    condition_for_pandas=""
                list_for_pandas.append([time_for_pandas,
                                        name_for_pandas,
                                        state_for_pandas,
                                        type_for_pandas,
                                        delay_for_pandas,
                                        condition_for_pandas,])
            return pandas.DataFrame(list_for_pandas,columns=["Time_report","Name_report", "ON_OFF_report","Type_report", "Delay_report", "Condition_report"])
        
        #словарь для #_report.lst, p1runout и p1spn
        self.data_dict={}
        
        #инициализация графиков, хронологии и ключевых параметров
        self.reset_data()
        
        #открываю папку RESULT
        os.chdir(path_to_folder)
        
        #считываю в формате DataFrame файл #_report.lst
        for report_name in os.listdir():
            if report_name.endswith("#_report.lst"):
                self.data_dict["report"]=parsing_report(report_name)

        #создаю и заполняю список из названий кластеров dia
        dia_names=[]
        for dia_name in os.listdir():
            if dia_name.endswith(".dia"):
                dia_names.append(dia_name)
        #создаю DataFrame из параметров кластеров dia
        dia_table=pandas.read_table(dia_names[0], skiprows=1, delim_whitespace=True)
        #создаю словарь, в котором ключи - названия кластеров dia, а значения - входящие в данный кластер параметры
        self.dia_clusters_dict={}
        self.dia_clusters_dict[dia_names[0].split("#")[-1].split(".dia")[0]]=list(dia_table.columns)

        #заполняю созданные DataFrame и словарь
        for dia_name in dia_names[1:]:
            table=pandas.read_table(dia_name, skiprows=1, delim_whitespace=True, encoding="mbcs")
            if len(table)==len(dia_table):
                dia_table=pandas.merge(dia_table,table)
                self.dia_clusters_dict[dia_name.split("#")[-1].split(".dia")[0]]=list(table.columns)
            else:
                # бывает, что СОКРАТ криво записывает файлы
                print("Проблемы с кластером " + dia_name)
        self.data_dict["dia"]=dia_table
        
        
        #считываю в формате DataFrame файлы p1runout и p1spn
        for folder_hef in os.listdir():
            if folder_hef.endswith("#hefest"):
                os.chdir(folder_hef)
                if "p1runout" in os.listdir():
                    self.data_dict["p1runout"]=pandas.read_table("p1runout", delim_whitespace=True)
                if "p1spn" in os.listdir():
                    self.data_dict["p1spn"]=pandas.read_table("p1spn", delim_whitespace=True)
        
        #вызов методов last_zero_parameters_transform и not_zero_parameters_transform при инициализации объекта
        self.last_zero_parameters_transform(names_last_zero_parameters)
        self.not_zero_parameters_transform(names_not_ziro_parameters)
        
        #Проверка разницы между Time и TIME_a
        if time_shift==None and "TIME_a" in dia_table:
            time_shift=dia_table["Time"][0] - dia_table["TIME_a"][0]
        
        if time_shift!=0 and time_shift!=None:
            self.time_shift_transform(time_shift)
        
        os.chdir(initial_dir)
        
    #метод для проверки параметров, которые если упали в 0, там и остаются (например температуры оболочек твэлов, СУЗов и топлива)
    def last_zero_parameters_transform(self,names_last_zero_parameters):
        if isinstance(names_last_zero_parameters,str):
            names_last_zero_parameters=[names_last_zero_parameters]
        for name_last_zero_parameter in names_last_zero_parameters:
            if name_last_zero_parameter in self.data_dict["dia"].columns and min(self.data_dict["dia"][name_last_zero_parameter])<=0:
                start_index=self.data_dict["dia"][self.data_dict["dia"][name_last_zero_parameter]<=0][name_last_zero_parameter].index[0]
                self.data_dict["dia"][name_last_zero_parameter][start_index:]=0
    
    #метод для проверки параметров, которые не должны резко падать в ноль
    def not_zero_parameters_transform(self,names_not_zero_parameters):
        if isinstance(names_not_zero_parameters,str):
            names_not_zero_parameters=[names_not_zero_parameters]
        for name_not_zero_parameter in names_not_zero_parameters:
            if name_not_zero_parameter in self.data_dict["dia"].columns and min(self.data_dict["dia"][name_not_zero_parameter])<=0:
                for k in range(len(self.data_dict["dia"][name_not_zero_parameter])-1):
                    if self.data_dict["dia"][name_not_zero_parameter][k+1]<=0:
                        self.data_dict["dia"][name_not_zero_parameter][k+1]=self.data_dict["dia"][name_not_zero_parameter][k]
        
    #метод для сдвига времени (прибавляет указанную величину к Time в кластерах dia, к Time в #_report.lst, к time в p1runout и к # в p1spn)
    def time_shift_transform(self, delta_time):
        for table_name in self.data_dict:
            self.data_dict[table_name][self.data_dict[table_name].columns[0]]+=delta_time
    
    #метод для обрезки значений снизу
    def cut_bottom_transform(self, parameter_name, parameter_value):
        time_row=self.get_row(parameter_name,parameter_value)
        keys_to_remote=[]
        if time_row is not None:
            lower_time=time_row[0]
            for table_name in self.data_dict:
                table=self.data_dict[table_name]
                time_table=table.iloc[0,0]
                if time_table<lower_time:
                    if len(table[table.columns[0]][table[table.columns[0]]>lower_time])>0:
                        lower_index=table[table.columns[0]][table[table.columns[0]]>lower_time].index[0]
                    else:
                        lower_index=len(table)
                    self.data_dict[table_name]=self.data_dict[table_name][lower_index:]
                    self.data_dict[table_name].reset_index(drop=True, inplace=True)
                    if len(self.data_dict[table_name])==0:
                        keys_to_remote.append(table_name)
        for key in keys_to_remote:
            self.data_dict.pop(key)
        
    #метод для обрезки значений сверху
    def cut_top_transform(self, parameter_name, parameter_value):
        time_row=self.get_row(parameter_name,parameter_value)
        keys_to_remote=[]
        if time_row is not None:
            upper_time=time_row[0]
            for table_name in self.data_dict:
                table=self.data_dict[table_name]
                time_table=table.iloc[-1,0]
                if time_table>upper_time:
                    if len(table[table.columns[0]][table[table.columns[0]]<upper_time])>0:
                        upper_index=table[table.columns[0]][table[table.columns[0]]>upper_time].index[0]
                    else:
                        upper_index=0
                    self.data_dict[table_name]=self.data_dict[table_name][:upper_index]
                    self.data_dict[table_name].reset_index(drop=True, inplace=True)
                    if len(self.data_dict[table_name])==0:
                        keys_to_remote.append(table_name)

        for key in keys_to_remote:
            self.data_dict.pop(key)

    #метод для получения информации о параметре
    def get_parameter_table (self,name_values):
        if isinstance(name_values,str):
            name_values=[name_values]
        return_table=None
        for table_name in self.data_dict:
            table=self.data_dict[table_name]
            if set(name_values) <= set(table.columns):
                return_table=table
        
        if return_table is not None:
            return return_table
        else:
            return None
    
    #метод для получения ряда
    def get_row (self, parameter_name, parameter_value, shift_down=True):
        table=self.get_parameter_table(parameter_name)
        if table is not None:
            parameter_min = min(table[parameter_name])
            parameter_max = max(table[parameter_name])
            if parameter_value == "min":
                parameter_value = parameter_min
            elif parameter_value == "max":
                parameter_value = parameter_max
        if table is not None and isinstance(parameter_value, float):
            if parameter_min<=parameter_value<=parameter_max:
                table_equal=table[parameter_name][table[parameter_name]==parameter_value]
                indexes=[]
                if len(table_equal)>0:
                    if shift_down:
                        indexes.append(table_equal.index[0])
                    else:
                        indexes.append(table_equal.index[-1])
                indexes_not_equal=[]
                for i in table[parameter_name].index[:-1]:
                    if table[parameter_name][i]<parameter_value<table[parameter_name][i+1] or table[parameter_name][i+1]<parameter_value<table[parameter_name][i]:
                        indexes_not_equal.append(i)

                if len(indexes_not_equal)>0:
                    if shift_down:
                        indexes.append(indexes_not_equal[0])
                    else:
                        indexes.append(indexes_not_equal[-1]+1)

                if shift_down:
                    return table.iloc[min(indexes),:]
                else:
                    return table.iloc[min(indexes),:]
            else:
                print("Значение {} не лежит в диапазоне параметра".format(parameter_name))
                return None
        else:
            print("Параметр {} не найден в dia, p1runout и p1spn".format(parameter_name))
            return None
    
    #метод для получения момента времени
    def get_time_moment(self, parameter_name, parameter_value, cut_down=None, cut_up=None):
        calc_copy=self.make_copy()
        return_value=None
        if isinstance(cut_down, dict):
            calc_copy.cut_bottom_transform(cut_down["name"], cut_down["value"])
        if isinstance(cut_up, dict):
            calc_copy.cut_top_transform(cut_up["name"], cut_up["value"])
        time_row=calc_copy.get_row(parameter_name, parameter_value)
        if time_row is not None:
            return_value=time_row[0]
        else:
            if "report" in calc_copy.data_dict:
                report_table=calc_copy.data_dict["report"][calc_copy.data_dict["report"]["Name_report"]==parameter_name]
                if len(report_table)>0:
                    if "ON" in parameter_value:
                        report_table=report_table[report_table["ON_OFF_report"]==True]
                    elif "OFF" in parameter_value:
                        report_table=report_table[report_table["ON_OFF_report"]==False]
                    else:
                        report_table=None
                    if report_table is not None:
                        if "LAST" in parameter_value:
                            return_value=report_table.iloc[-1,0]
                        elif "FIRST" in parameter_value:
                            return_value=report_table.iloc[0,0]
        return return_value
        
    #метод для получения ключевого параметра
    def get_key_parameter(self, key_parameter_name, search_name, search_value, cut_down=None, cut_up=None):
        time_moment=self.get_time_moment(search_name, search_value, cut_down, cut_up)
        key_table=self.get_parameter_table(key_parameter_name)
        return_value=None
        if time_moment is not None and key_table is not None:
            key_table=key_table[key_table.iloc[:,0]<=time_moment]
            if len(key_table)>0:
                key_row=key_table.iloc[-1]
                return_value=key_row[key_parameter_name]
        return return_value
        
    
    
    #метод для получения значений
    def get_values(self,name_values):
        parameter_table=self.get_parameter_table(name_values)
        if parameter_table is not None:
            return parameter_table[name_values]
        else:
            print("Имя " + name_values + " не найдено")
            return None
    #метод для получения dia_кластера
    def get_dia_cluster(self, cluster_name):
        if cluster_name in self.dia_clusters_dict:
            return self.data_dict["dia"][self.dia_clusters_dict[cluster_name]]
        else:
            return None
    #метод для умножения и смещения значений параметра
    def values_mult_shift(self, name_parameter, mult=None, shift=None):
        if mult is None:
            mult=1
        if shift is None:
            shift=0
        for table_name in self.data_dict:
            if name_parameter in self.data_dict[table_name]:
                self.data_dict[table_name][name_parameter]=self.data_dict[table_name][name_parameter]*mult+shift
            
    #метод для копирования объекта
    def make_copy (self):
        import copy
        return copy.deepcopy(self)
    
    #метод для отрисовки графика
    def graph (self, GrName, x_names, y_names, lablex="t, s", labley="", x1=None, y1=None, x2=None, y2=None, mult_x=None, mult_y=None, shift_x=None, shift_y=None, stpx=None, stpy=None):
        
        if isinstance(y_names, str):
            y_names=[y_names]
        if isinstance(x_names, str):
            x_names=[x_names for i in range(len(y_names))]
        
        if len(y_names)!=len(x_names):
            x_names=[x_names[0] for i in range(len(y_names))]
        
        
        
        get_x=[self.get_values(x_name) for x_name in x_names]
        get_y=[self.get_values(y_name) for y_name in y_names]
        
        path_dict=graph(GrName, get_x, get_y, lablex=lablex, labley=labley, x1=x1, y1=y1, x2=x2, y2=y2, mult_x=mult_x, mult_y=mult_y, shift_x=shift_x, shift_y=shift_y, stpx=stpx, stpy=stpy)
        
        path_dict["x"]="_".join(x_names)
        path_dict["y"]="_".join(y_names)
        self.graph_list.append(path_dict)
    
    #метод для создания хронологии
    def make_chrono(self,time_list):
        import pandas
        if isinstance(time_list,dict):
            time_list=[time_list]
        chrono_list=[self.chrono_table]
        for time_point in time_list:
            time_moment=self.get_time_moment(time_point["name"],time_point["value"],time_point["cut_down"],time_point["cut_up"])
            if time_moment is not None:
                chrono_list.append(pandas.DataFrame([[time_moment,time_point["description"]]],columns=["Time","Description"]))
        self.chrono_table=pandas.concat(chrono_list).drop_duplicates().sort_values(by=["Time"]).reset_index(drop=True)
    
    #метод для создания DataFrame с ключевыми параметрами
    def make_key_table(self,key_list):
        import pandas
        if isinstance(key_list,dict):
            key_list=[key_list]
        key_parameters_list=[self.key_parameters_table]
        for key_point in key_list:
            key_value=self.get_key_parameter(key_point["key_parameter_name"],key_point["search_name"],key_point["value"],key_point["cut_down"],key_point["cut_up"])
            if key_value is not None:
                 key_parameters_list.append(pandas.DataFrame([[key_point["description"],key_value]],columns=["Description","Value"]))
        self.key_parameters_table=pandas.concat(key_parameters_list).drop_duplicates().reset_index(drop=True)
    
    #метод для сброса ключевых параметров, хронологий и графиков
    def reset_data(self):
        import pandas
        self.graph_list=[]
        self.chrono_table=pandas.DataFrame(columns=["Time","Description"])
        self.key_parameters_table=pandas.DataFrame(columns=["Description","Value"])
    
    
    #метод для сохранения в файл
    def save(self, path="", name="data"):
        import pickle
        copy_to_save=self.make_copy()
        copy_to_save.reset_data()
        with open(path+name+'.sokle', 'wb') as f:
            pickle.dump(copy_to_save, f)

class socrat_calculation_series:
    def __init__(self, calculations):
        import pandas
        self.series=[]
        self.graph_list=[]
        self.chrono_table=pandas.DataFrame(columns=["Description"])
        self.key_parameters_table=pandas.DataFrame(columns=["Description"])
        if not isinstance(calculations,list):
            calculations=[calculations]
        for i in range(len(calculations)):
            self.series.append(calculations[i])
            calc_chrono=calculations[i].chrono_table.reindex(columns=["Description", "Time"])
            calc_chrono.columns=["Description", "Time_"+str(i+1)]
            self.chrono_table=self.chrono_table.merge(calc_chrono,how="outer",on = "Description").fillna("–")
            calc_key_parameters=calculations[i].key_parameters_table.copy()
            calc_key_parameters.columns=["Description", "Values_"+str(i+1)]
            self.key_parameters_table=self.key_parameters_table.merge(calc_key_parameters,how="outer",on = "Description").fillna("–")

    
    def graph (self, GrName, x_names, y_names, lablex="t, s", labley="", x1=None, y1=None, x2=None, y2=None, mult_x=None, mult_y=None, shift_x=None, shift_y=None, stpx=None, stpy=None):
        
        if isinstance(y_names, str):
            y_names=[y_names]
        if isinstance(x_names, str):
            x_names=[x_names for i in range(len(y_names))]
        
        get_x=[]
        get_y=[]
        
        for calc in self.series:
            for i in range(len(y_names)):
                get_x.append(calc.get_values(x_names[i]))
                get_y.append(calc.get_values(y_names[i]))
        
        path_dict=graph(GrName, get_x, get_y, lablex=lablex, labley=labley, x1=x1, y1=y1, x2=x2, y2=y2, mult_x=mult_x, mult_y=mult_y, shift_x=shift_x, shift_y=shift_y, stpx=stpx, stpy=stpy)
        
        path_dict["x"]="_".join(x_names)
        path_dict["y"]="_".join(y_names)
        self.graph_list.append(path_dict)




if __name__ == '__main__':

    import pandas
    import os
    import docx
    
    name_input_excel="Pre_POOL.xlsx"

    excel_sheets=list(pandas.read_excel(name_input_excel,sheet_name=None, header=None).values())
    excel_sheets=[excel_input(i) for i in excel_sheets]

    for excel_sheet in excel_sheets:
        if excel_sheet.results_src is not None:
            for i in range(len(excel_sheet.results_src)):
                calc=socrat_calculation(excel_sheet.results_src[i])
                if excel_sheet.last_zero_parameters is not None:
                    calc.last_zero_parameters_transform(excel_sheet.last_zero_parameters)
                if excel_sheet.not_zero_parameters is not None:
                    calc.not_zero_parameters_transform(excel_sheet.not_zero_parameters)
                if excel_sheet.cut_bottom_parameters is not None:
                    calc.cut_bottom_transform(excel_sheet.cut_bottom_parameters["name"], excel_sheet.cut_bottom_parameters["value"])
                if excel_sheet.cut_top_parameters is not None:
                    calc.cut_top_transform(excel_sheet.cut_top_parameters["name"], excel_sheet.cut_top_parameters["value"])

                if not os.path.exists(excel_sheet.results_post[i]):
                    os.mkdir(excel_sheet.results_post[i])
                os.chdir(excel_sheet.results_post[i])
                calc.save(excel_sheet.results_post[i])
        else:
            import pickle
            calculations_list=[]
            for i in range(len(excel_sheet.results_post)):
                with open(excel_sheet.results_post[i]+'data.sokle', 'rb') as f:
                    calc=pickle.load(f)

                    if excel_sheet.last_zero_parameters is not None:
                        calc.last_zero_parameters_transform(excel_sheet.last_zero_parameters)
                    if excel_sheet.not_zero_parameters is not None:
                        calc.not_zero_parameters_transform(excel_sheet.not_zero_parameters)
                    if excel_sheet.cut_bottom_parameters is not None:
                        calc.cut_bottom_transform(excel_sheet.cut_bottom_parameters["name"], excel_sheet.cut_bottom_parameters["value"])
                    if excel_sheet.cut_top_parameters is not None:
                        calc.cut_top_transform(excel_sheet.cut_top_parameters["name"], excel_sheet.cut_top_parameters["value"])
                    if excel_sheet.corrected_parameters is not None:
                        for corrected_parameter in excel_sheet.corrected_parameters:
                            calc.values_mult_shift(corrected_parameter, mult=excel_sheet.corrected_parameters[corrected_parameter]["mult"], shift=excel_sheet.corrected_parameters[corrected_parameter]["shift"])
                    if excel_sheet.time_parameters is not None:
                        calc.make_chrono(excel_sheet.time_parameters)

                    if excel_sheet.key_parameters is not None:
                        calc.make_key_table(excel_sheet.key_parameters, P_CORE_OUT_name=excel_sheet.P_CORE_OUT_name)

                    calculations_list.append(calc)
            if len(excel_sheet.results_post)==len(excel_sheet.results_final):
                for calculation_number in range(len(excel_sheet.results_post)):

                    if not os.path.exists(excel_sheet.results_final[calculation_number]):
                        os.mkdir(excel_sheet.results_final[calculation_number])
                    os.chdir(excel_sheet.results_final[calculation_number])

                    if excel_sheet.critic_clusters is not None:
                        critic(calculations_list[calculation_number],excel_sheet.critic_clusters,excel_sheet.results_final[calculation_number])

                    if excel_sheet.graphs is not None:
                        for graph_number in range(len(excel_sheet.graphs)):
                            calculations_list[calculation_number].graph(str(graph_number+1),
                                       excel_sheet.graphs[graph_number]["x"],
                                       excel_sheet.graphs[graph_number]["y"],
                                       lablex=excel_sheet.graphs[graph_number]["label_x"],
                                       labley=excel_sheet.graphs[graph_number]["label_y"],
                                       x1=excel_sheet.graphs[graph_number]["min_x"],
                                       x2=excel_sheet.graphs[graph_number]["max_x"],
                                       y1=excel_sheet.graphs[graph_number]["min_y"],
                                       y2=excel_sheet.graphs[graph_number]["max_y"],
                                       mult_x=excel_sheet.graphs[graph_number]["mult_x"],
                                       mult_y=excel_sheet.graphs[graph_number]["mult_y"],
                                       stpx=excel_sheet.graphs[graph_number]["step_x"],
                                       stpy=excel_sheet.graphs[graph_number]["step_y"])

                    doc=docx.Document()
                    if len(calculations_list[calculation_number].chrono_table)>0:
                        doc.add_paragraph("Хронология")
                        table=doc.add_table(rows=len(calculations_list[calculation_number].chrono_table)+1, cols=2)
                        table.style="Table Grid"
                        table.cell(0,0).text="Время, с"
                        table.cell(0,1).text="Событие"
                        for k in range(len(calculations_list[calculation_number].chrono_table)):
                            table.cell(k+1,0).text=str(calculations_list[calculation_number].chrono_table.iloc[k,0])
                            table.cell(k+1,1).text=str(calculations_list[calculation_number].chrono_table.iloc[k,1])
                        doc.add_page_break()

                    if len(calculations_list[calculation_number].key_parameters_table)>0:
                        doc.add_paragraph("Ключевые параметры")
                        table=doc.add_table(rows=1, cols=2)
                        table.style="Table Grid"
                        table.cell(0,0).text="Параметр"
                        table.cell(0,1).text="Значение"
                        for k in range(len(calculations_list[calculation_number].key_parameters_table)):
                            table.add_row()
                            table.cell(k+1,0).text=str(calculations_list[calculation_number].key_parameters_table.iloc[k,0])
                            table.cell(k+1,1).text=str(calculations_list[calculation_number].key_parameters_table.iloc[k,1])
                        doc.add_page_break()

                    for k in range(len(calculations_list[calculation_number].graph_list)):
                        doc.add_picture(calculations_list[calculation_number].graph_list[k]["png"], width = docx.shared.Cm(12))
                        doc.add_paragraph("Рисунок {}\n{}".format(str(k+1),excel_sheet.graphs[k]["description"]))
                    doc.save("Results.docx")
            else:
                if not os.path.exists(excel_sheet.results_final[0]):
                    os.mkdir(excel_sheet.results_final[0])
                os.chdir(excel_sheet.results_final[0])
                calculations=socrat_calculation_series(calculations_list)
                if excel_sheet.graphs is not None:
                    for graph_number in range(len(excel_sheet.graphs)):
                        calculations.graph(str(graph_number+1),
                                   excel_sheet.graphs[graph_number]["x"],
                                   excel_sheet.graphs[graph_number]["y"],
                                   lablex=excel_sheet.graphs[graph_number]["label_x"],
                                   labley=excel_sheet.graphs[graph_number]["label_y"],
                                   x1=excel_sheet.graphs[graph_number]["min_x"],
                                   x2=excel_sheet.graphs[graph_number]["max_x"],
                                   y1=excel_sheet.graphs[graph_number]["min_y"],
                                   y2=excel_sheet.graphs[graph_number]["max_y"],
                                   mult_x=excel_sheet.graphs[graph_number]["mult_x"],
                                   mult_y=excel_sheet.graphs[graph_number]["mult_y"],
                                   stpx=excel_sheet.graphs[graph_number]["step_x"],
                                   stpy=excel_sheet.graphs[graph_number]["step_y"])
                doc=docx.Document()
                if len(calculations.chrono_table)>0:
                    doc.add_paragraph("Хронология")
                    table=doc.add_table(rows=len(calculations.chrono_table)+1, cols=len(calculations.chrono_table.columns))
                    table.style="Table Grid"
                    for i in range(len(calculations.chrono_table.columns)):
                        table.cell(0,i).text=calculations.chrono_table.columns[i]
                        for k in range(len(calculations.chrono_table)):
                            table.cell(k+1,i).text=str(calculations.chrono_table.iloc[k,i])
                    table.cell(0,0).text="Событие"
                    doc.add_page_break()

                if len(calculations.key_parameters_table)>0:
                    doc.add_paragraph("Ключевые параметры")
                    table=doc.add_table(rows=len(calculations.key_parameters_table)+1, cols=len(calculations.key_parameters_table.columns))
                    table.style="Table Grid"
                    for i in range(len(calculations.key_parameters_table.columns)):
                        table.cell(0,i).text=calculations.key_parameters_table.columns[i]
                        for k in range(len(calculations.key_parameters_table)):
                            table.cell(k+1,i).text=str(calculations.key_parameters_table.iloc[k,i])
                    table.cell(0,0).text="Параметр"
                    doc.add_page_break()

                for k in range(len(calculations.graph_list)):
                    doc.add_picture(calculations.graph_list[k]["png"], width = docx.shared.Cm(12))
                    doc.add_paragraph("Рисунок {}\n{}".format(str(k+1),excel_sheet.graphs[k]["description"]))

                doc.save("Results.docx")


