str_help = """Данная программа предназначена для построения хронологий, таблиц с ключевыми параметрами и графиков"""


import sys
import os
import pickle
import traceback
import docx
import json

from PySide6.QtWidgets import QFileDialog, QDialog
from PySide6.QtCore import Qt, QThreadPool, QObject, Signal, Slot, QRunnable
from PySide6.QtWidgets import QApplication, QMainWindow, QTreeWidgetItem, QCompleter, QMessageBox, QMenu
from PySide6.QtGui import QAction, QTextCursor

from ui_main_window import Ui_MainWindow
from socrat_eye import Socrat_calculation, Trap_calculation, Series_of_calculations
from graphs_edit_window import Ui_Dialog_graphs_edit
from chrono_edit_window import Ui_Dialog_chrono_edit
from key_parameters_edit_window import Ui_Dialog_key_parameters_edit
from correcting_edit_window import Ui_Dialog_correcting_edit

class MainWindow(QMainWindow):
    def __init__(self, parent=None):
        super(MainWindow, self).__init__(parent)
        self.ui = Ui_MainWindow()
        self.ui.setupUi(self)

        # Создание менеджера пула
        self.thread_manager = QThreadPool()
        self.ui.button_stop.clicked.connect(self.stop_run)
        # Перенаправление потока вывода
        stdout = OutputWrapper(self, True)
        stdout.outputWritten.connect(self.handleOutput)
        stderr = OutputWrapper(self, False)
        stderr.outputWritten.connect(self.handleOutput)

        # Справка
        print(str_help)

        # Инициализация списка с расчётами
        self.calculations = []

        # Инициализация буфера расчетов
        self.contextData = []
        # Инициализация буфера с корректируемыми параметрами
        self.contextCorrecting = []
        # Инициализация буфера графиков
        self.contextGraphs = []
        # Инициализация буфера хронологий
        self.contextChrono = []
        # Инициализация буфера ключевых параметров
        self.contextKeyParameters = []

        # Инициализация completer
        self.completer = Completers()
        # Подключение полей к completer для автодополнения
        # Глобальная обрезка снизу
        self.ui.line_cut_down_name.setCompleter(self.completer.clusters_completer)
        self.ui.line_cut_down_value.setCompleter(self.completer.help_clusters_completer)
        # Глобальная обрезка сверху
        self.ui.line_cut_up_name.setCompleter(self.completer.clusters_completer)
        self.ui.line_cut_up_value.setCompleter(self.completer.help_clusters_completer)
        # Ключевые параметры
        self.ui.line_key_parameter_name.setCompleter(self.completer.clusters_completer)
        self.ui.line_search_parameter_name.textChanged.connect(lambda:
                                                               self.selection_completer_clusters_report(
                                                                   self.ui.line_search_parameter_name,
                                                                   self.ui.line_search_parameter_value))
        self.ui.line_search_parameter_value.textChanged.connect(lambda:
                                                                self.selection_completer_help(
                                                                    self.ui.line_search_parameter_name,
                                                                    self.ui.line_search_parameter_value))
        # Обрезка снизу ключевых параметров
        self.ui.line_cut_down_key_parameters_name.setCompleter(self.completer.clusters_completer)
        self.ui.line_cut_down_key_parameters_value.setCompleter(self.completer.help_clusters_completer)
        # Обрезка сверху ключевых параметров
        self.ui.line_cut_up_key_parameters_name.setCompleter(self.completer.clusters_completer)
        self.ui.line_cut_up_key_parameters_value.setCompleter(self.completer.help_clusters_completer)
        # Хронология
        self.ui.line_chrono_parameter_name.textChanged.connect(lambda:
                                                               self.selection_completer_clusters_report(
                                                                   self.ui.line_chrono_parameter_name,
                                                                   self.ui.line_chrono_parameter_value))
        self.ui.line_chrono_parameter_value.textChanged.connect(lambda:
                                                                self.selection_completer_help(
                                                                    self.ui.line_chrono_parameter_name,
                                                                    self.ui.line_chrono_parameter_value))
        # Обрезка хронологий снизу
        self.ui.line_cut_down_chrono_name.setCompleter(self.completer.clusters_completer)
        self.ui.line_cut_down_chrono_value.setCompleter(self.completer.help_clusters_completer)
        # Обрезка хронологий сверху
        self.ui.line_cut_up_chrono_name.setCompleter(self.completer.clusters_completer)
        self.ui.line_cut_up_chrono_value.setCompleter(self.completer.help_clusters_completer)

        # Корректируемые параметры
        # связь с completer
        self.ui.line_corrected_parameters_name.setCompleter(self.completer.clusters_completer)
        # методы для обновления completer при изменении корректируемых параметров
        self.ui.line_corrected_parameters_name.textChanged.connect(
            lambda: self.completer.update_correcting_parameters(self.ui.treeWidget_corrected_parameters))
        self.ui.tabWidget_main.currentChanged.connect(
            lambda: self.completer.update_correcting_parameters(self.ui.treeWidget_corrected_parameters))
        # Графики
        self.ui.line_graphs_x_names.setCompleter(self.completer.graphs_completer)
        self.ui.line_graphs_x_names.textChanged.connect(
            lambda: self.completer.update_graphs(self.ui.line_graphs_y_names))

        self.ui.line_graphs_y_names.setCompleter(self.completer.graphs_completer)
        self.ui.line_graphs_y_names.textChanged.connect(
            lambda: self.completer.update_graphs(self.ui.line_graphs_y_names))


        # Подключение общих кнопок
        # Кнопка для выбора пути сохранения результатов
        self.ui.button_path_save.clicked.connect(
            lambda: self.ui.line_path_to_save.setText(QFileDialog.getExistingDirectory()))
        # Кнопка для предзагрузки
        self.ui.button_pre_run.clicked.connect(self.pre_run)
        # Кнопка выполнить
        self.ui.button_execute.clicked.connect(lambda: self.run_function_safely(self.run_programm))

        # Вкладка Данные
        # подключение кнопок для открытия файлов и папок
        # Кнопка для выбора пути с сырыми результатами
        self.ui.button_path_to_calc.clicked.connect(
            lambda: self.ui.line_path_to_result.setText(QFileDialog.getExistingDirectory()))
        # Кнопка для выбора пути с бинарными результатами
        self.ui.button_path_to_sackle.clicked.connect(self.open_sackle_file)
        # подключение кнопок для добавления расчётов
        self.ui.button_add_sackle.clicked.connect(self.add_sockle_file)
        self.ui.button_add_result.clicked.connect(self.add_result_folder)
        # Подключение кнопок для сброса и удаления расчётов
        self.ui.button_reset_data.clicked.connect(lambda: self.ui.treeWidget_calculations.clear())
        self.ui.button_del_chosen_data.clicked.connect(
            lambda: self.del_selected_elements(self.ui.treeWidget_calculations))
        # Перемещение вверх и вниз
        self.ui.button_up_data.clicked.connect(
            lambda: self.move_selected_up(self.ui.treeWidget_calculations))
        self.ui.button_down_data.clicked.connect(
            lambda: self.move_selected_down(self.ui.treeWidget_calculations))
        # Подключение контекстного меню
        self.ui.treeWidget_calculations.customContextMenuRequested.connect(
            lambda position: self.open_context_menu(position,
                                                    self.ui.treeWidget_calculations,
                                                    self.contextData))
        # Вкладка Обработка
        # Подключение кнопок для добавления параметров
        self.ui.button_corrected_parameters_add.clicked.connect(self.add_corected_parameter)
        # Подключение кнопок для сброса и удаления параметров
        self.ui.button_corrected_parameters_reset.clicked.connect(
            lambda: self.ui.treeWidget_corrected_parameters.clear())
        self.ui.button_corrected_parameters_del.clicked.connect(
            lambda: self.del_selected_elements(self.ui.treeWidget_corrected_parameters))

        # Импорт и экспорт настроек
        self.ui.button_export_cfg.clicked.connect(self.save_cfg)
        self.ui.button_import_cfg.clicked.connect(self.open_sfg)
        # Редактирование при двойном клике
        self.ui.treeWidget_corrected_parameters.itemDoubleClicked.connect(self.edit_correcting)

        # Перемещение вверх и вниз
        self.ui.button_up_correcting.clicked.connect(lambda : self.move_selected_up(self.ui.treeWidget_corrected_parameters))
        self.ui.button_down_correcting.clicked.connect(lambda: self.move_selected_down(self.ui.treeWidget_corrected_parameters))

        # Подключение контекстного меню
        self.ui.treeWidget_corrected_parameters.customContextMenuRequested.connect(
            lambda position: self.open_context_menu(position,
                                                    self.ui.treeWidget_corrected_parameters,
                                                    self.contextCorrecting))

        # Вкладка Графики
        # Подключение кнопок для добавления графиков
        self.ui.button_graphs_add.clicked.connect(self.add_graph)
        # Подключение кнопок для сброса и удаления графиков
        self.ui.button_reset_graphs.clicked.connect(lambda: self.get_tree_graphs().clear())
        self.ui.button_del_chosen_graphs.clicked.connect(self.del_selected_graphs)
        # Подключение кнопок для перемещения графиков
        self.ui.button_up_graphs.clicked.connect(self.move_up_selected_graphs)
        self.ui.button_down_graphs.clicked.connect(self.move_down_selected_graphs)
        # Редактирование при двойном клике
        self.ui.treeWidget_graphs_single.itemDoubleClicked.connect(self.edit_graph)
        self.ui.treeWidget_graphs_multy.itemDoubleClicked.connect(self.edit_graph)
        # Подключение контекстного меню
        self.ui.treeWidget_graphs_single.customContextMenuRequested.connect(
            lambda position: self.open_context_menu(position,
                                                    self.ui.treeWidget_graphs_single,
                                                    self.contextGraphs))
        self.ui.treeWidget_graphs_multy.customContextMenuRequested.connect(
            lambda position: self.open_context_menu(position,
                                                    self.ui.treeWidget_graphs_multy,
                                                    self.contextGraphs))

        # Вкладка Хронология
        # Подключение кнопок для добавления события
        self.ui.button_chrono_parameters_add.clicked.connect(self.add_chrono)
        # Подключение кнопок для сброса и удаления событий
        self.ui.button_reset_chrono_parameters.clicked.connect(lambda: self.get_tree_chrono().clear())
        self.ui.button_del_chosen_chrono_parameters.clicked.connect(
            lambda: self.del_selected_elements(self.get_tree_chrono()))
        # Редактирование при двойном клике
        self.ui.treeWidget_chrono_single.itemDoubleClicked.connect(self.edit_chrono)
        self.ui.treeWidget_chrono_multy.itemDoubleClicked.connect(self.edit_chrono)
        # Подключение контекстного меню
        self.ui.treeWidget_chrono_single.customContextMenuRequested.connect(
            lambda position: self.open_context_menu(position,
                                                    self.ui.treeWidget_chrono_single,
                                                    self.contextChrono))
        self.ui.treeWidget_chrono_multy.customContextMenuRequested.connect(
            lambda position: self.open_context_menu(position,
                                                    self.ui.treeWidget_chrono_multy,
                                                    self.contextChrono))

        # Вкладка Ключевые параметры
        # Подключение кнопок для добавления ключевых параметров
        self.ui.button_key_parameters_add.clicked.connect(self.add_key_parameter)
        # Подключение кнопок для сброса и удаления ключевых параметров
        self.ui.button_reset_key_parameters.clicked.connect(lambda: self.get_tree_key_parameters().clear())
        self.ui.button_del_chosen_key_parameters.clicked.connect(
            lambda: self.del_selected_elements(self.get_tree_key_parameters()))
        # Подключение кнопок для перемещения ключевых параметров
        self.ui.button_up_key_parameters.clicked.connect(self.move_up_selected_key_parameters)
        self.ui.button_down_key_parameters.clicked.connect(self.move_down_selected_key_parameters)
        # Редактирование при двойном клике
        self.ui.treeWidget_key_parameters_single.itemDoubleClicked.connect(self.edit_key_parameter)
        self.ui.treeWidget_key_parameters_multy.itemDoubleClicked.connect(self.edit_key_parameter)
        # Подключение контекстного меню
        self.ui.treeWidget_key_parameters_single.customContextMenuRequested.connect(
            lambda position: self.open_context_menu(position,
                                                    self.ui.treeWidget_key_parameters_single,
                                                    self.contextKeyParameters))
        self.ui.treeWidget_key_parameters_multy.customContextMenuRequested.connect(
            lambda position: self.open_context_menu(position,
                                                    self.ui.treeWidget_key_parameters_multy,
                                                    self.contextKeyParameters))

    # метод для перенаправления вывода
    def handleOutput(self, text, stdout):
        self.ui.textBrowser_terminal.moveCursor(QTextCursor.End)
        self.ui.textBrowser_terminal.setTextColor(Qt.black if stdout else Qt.red)
        self.ui.textBrowser_terminal.insertPlainText(text)
        self.ui.textBrowser_terminal.setTextColor(Qt.black)

    # методы для автодополнений
    # метод, который будет возвращать completer в поле с названием в зависимости от поля значений
    def selection_completer_clusters_report(self, line_name, line_value):
        value_str = line_value.text()
        if value_str in ("min", "max"):
            line_name.setCompleter(self.completer.clusters_completer)
        elif value_str in ("ON_FIRST", "ON_LAST", "OFF_FIRST", "OFF_LAST"):
            line_name.setCompleter(self.completer.report_completer)
        else:
            line_name.setCompleter(self.completer.clusters_report_completer)

    # метод для выбора completer в поле со значением в зависимости от поля с названием
    def selection_completer_help(self, line_name, line_value):
        name_str = line_name.text()
        if name_str in self.completer.set_clusters:
            line_value.setCompleter(self.completer.help_clusters_completer)
        elif name_str in self.completer.set_report:
            line_value.setCompleter(self.completer.help_report_completer)
        else:
            line_value.setCompleter(self.completer.help_clusters_report_completer)

    # метод для отображения окна с предупреждением
    def show_message_window(self, main_text, informative_text=None):
        msgBox = QMessageBox()
        msgBox.setText(main_text)
        if isinstance(informative_text,str):
            msgBox.setInformativeText(informative_text)
        msgBox.setWindowIcon(self.windowIcon())
        msgBox.setWindowTitle("Сообщение")
        msgBox.exec()

    # методы для общих кнопок
    # метод для удаления выбранных полей
    def del_selected_elements (self, tree):
        selected_items = tree.selectedItems()
        for item in selected_items:
            index = tree.indexOfTopLevelItem(item)
            tree.takeTopLevelItem(index)

    def past_selected_elements (self, tree, buffer):
        selected_items = tree.selectedItems()
        selected_indexes = [tree.indexOfTopLevelItem(item) for item in selected_items]
        if not selected_indexes:
            selected_indexes = [tree.topLevelItemCount()]
        index = selected_indexes[-1]
        if buffer:
            if buffer[0].treeWidget() is not tree:
                buffer[:] = [item.clone() for item in buffer]
            tree.insertTopLevelItems(index, tuple(buffer))
        self.reset_index_graphs()

    def copy_selected_elements (self, tree, buffer):
        buffer[:] = tree.selectedItems()

    def cut_selected_elements(self, tree, buffer):
        self.copy_selected_elements(tree,buffer)
        self.del_selected_elements(tree)
        self.reset_index_graphs()
    def select_all_elements(self, tree):
        for i in range(tree.topLevelItemCount()): tree.topLevelItem(i).setSelected(True)

    # метод для перемещения элемента вверх
    def move_selected_up(self, tree):
        selected_items = tree.selectedItems()
        selected_indexes = [tree.indexOfTopLevelItem(item) for item in selected_items]
        if selected_indexes and min(selected_indexes)>0:
            tree.clearSelection()
            for i in selected_indexes:
                item = tree.takeTopLevelItem(i)
                tree.insertTopLevelItem(i-1, item)
            for item in selected_items:
                item.setSelected(True)

    # метод для перемещения элемента вниз
    def move_selected_down(self, tree):
        selected_items = tree.selectedItems()
        selected_indexes = [tree.indexOfTopLevelItem(item) for item in selected_items]
        if selected_indexes and max(selected_indexes) < tree.topLevelItemCount() - 1:
            tree.clearSelection()
            for i in selected_indexes[::-1]:
                item = tree.takeTopLevelItem(i)
                tree.insertTopLevelItem(i+1, item)
            for item in selected_items:
                item.setSelected(True)


    # метод для вызова контекстного меню
    def open_context_menu(self, position, tree, buffer):
        menu = QMenu()
        menu_cut = QAction("Вырезать", self)
        menu_copy = QAction("Копировать", self)
        menu_past = QAction("Вставить", self)
        menu_select_all = QAction("Выделить все", self)

        menu.addActions((menu_cut,menu_copy,menu_past,menu_select_all))
        menu_cut.triggered.connect(lambda: self.cut_selected_elements(tree, buffer))
        menu_copy.triggered.connect(lambda: self.copy_selected_elements(tree, buffer))
        menu_past.triggered.connect(lambda: self.past_selected_elements(tree, buffer))
        menu_select_all.triggered.connect(lambda: self.select_all_elements(tree))

        menu.exec(tree.viewport().mapToGlobal(position))





    # Вкладка Данные
    # методы для открытия файлов и папок
    def open_sackle_file(self):
        paths = QFileDialog.getOpenFileNames(self, "Открыть файл", "", "Файлы sockle (*.sokle)")[0]
        if len(paths) > 0:
            self.ui.line_path_to_sackle.setText(";".join(paths))
            self.add_sockle_file()

    # методы для добавления файлов и папок
    def add_calc(self, name, path):
        if len(self.ui.treeWidget_calculations.findItems(name, Qt.MatchFlag.MatchExactly)) == 0:
            item = QTreeWidgetItem([name,
                                    ";".join([self.ui.line_cut_down_name.text(), self.ui.line_cut_down_value.text()]),
                                    ";".join([self.ui.line_cut_up_name.text(), self.ui.line_cut_up_value.text()]),
                                    path, ])
            self.ui.treeWidget_calculations.addTopLevelItem(item)

            self.ui.line_path_to_sackle.setText("")
            self.ui.line_path_to_result.setText("")
            self.ui.line_name_calc.setText("")
        else:
            print("Такое имя уже существует")
            self.show_message_window("Такое имя уже существует")

    def add_sockle_file(self):
        paths = self.ui.line_path_to_sackle.text()
        paths = paths.split(";")
        for path in paths:
            if os.path.exists(path):
                name = path.split(r"/")[-1]
                if name.endswith(".sokle"):
                    name = name.split(".sokle")[-2]
                    self.add_calc(name, path)
                else:
                    print("Неверно задан путь")
                    self.show_message_window("Неверно задан путь")

    def add_result_folder(self):
        path = self.ui.line_path_to_result.text()
        name = self.ui.line_name_calc.text()
        if name != "":
            if os.path.exists(path):
                self.add_calc(name, path)
            else:
                print("Данного пути не существует")
                self.show_message_window("Данного пути не существует")
        else:
            print("Введено пустое имя")
            self.show_message_window("Введено пустое имя")

    # Вкладка Обработка
    # метод для добавления параметров
    def add_corected_parameter(self):
        name = self.ui.line_corrected_parameters_name.text()
        extra_operations=""
        if self.ui.radioButton_gradient.isChecked():
            extra_operations = self.ui.radioButton_gradient.text()
        elif self.ui.radioButton_integral.isChecked():
            extra_operations = self.ui.radioButton_integral.text()
        expression = self.ui.text_edit_correcting_expression.toPlainText()
        if name != "":
            if expression!="":
                item = QTreeWidgetItem([name, extra_operations, expression])
                self.ui.treeWidget_corrected_parameters.addTopLevelItem(item)
                # Нужно для completer
                self.ui.treeWidget_corrected_parameters.setCurrentItem(item)

                self.ui.line_corrected_parameters_name.setText("")
                self.ui.text_edit_correcting_expression.setText("")
            else:
                print("Введено пустое выражение")
                self.show_message_window("Введено пустое выражение")
        else:
            print("Введено пустое имя")
            self.show_message_window("Введено пустое имя")

    def save_cfg(self):
        config_dict={}
        # Вкладка Обработка
        config_dict["Correcting"]={}

        config_dict["Correcting"]["Items"]=[]
        for i in range(self.ui.treeWidget_corrected_parameters.topLevelItemCount()):
            item = self.ui.treeWidget_corrected_parameters.topLevelItem(i)
            config_dict["Correcting"]["Items"].append([])
            for k in range(item.columnCount()):
                config_dict["Correcting"]["Items"][-1].append(item.text(k))

        # Вкладка Графики
        config_dict["Graphs_single"]=[]
        for i in range(self.ui.treeWidget_graphs_single.topLevelItemCount()):
            item = self.ui.treeWidget_graphs_single.topLevelItem(i)
            config_dict["Graphs_single"].append([])
            for k in range(item.columnCount()):
                config_dict["Graphs_single"][-1].append(item.text(k))

        config_dict["Graphs_multy"] = []
        for i in range(self.ui.treeWidget_graphs_multy.topLevelItemCount()):
            item = self.ui.treeWidget_graphs_multy.topLevelItem(i)
            config_dict["Graphs_multy"].append([])
            for k in range(item.columnCount()):
                config_dict["Graphs_multy"][-1].append(item.text(k))

        # Вкладка Хронология
        config_dict["Chrono_single"] = []
        for i in range(self.ui.treeWidget_chrono_single.topLevelItemCount()):
            item = self.ui.treeWidget_chrono_single.topLevelItem(i)
            config_dict["Chrono_single"].append([])
            for k in range(item.columnCount()):
                config_dict["Chrono_single"][-1].append(item.text(k))
        config_dict["Chrono_multy"] = []
        for i in range(self.ui.treeWidget_chrono_multy.topLevelItemCount()):
            item = self.ui.treeWidget_chrono_multy.topLevelItem(i)
            config_dict["Chrono_multy"].append([])
            for k in range(item.columnCount()):
                config_dict["Chrono_multy"][-1].append(item.text(k))

        # Вкладка Ключевые параметры
        config_dict["Key_parameters_single"] = []
        for i in range(self.ui.treeWidget_key_parameters_single.topLevelItemCount()):
            item = self.ui.treeWidget_key_parameters_single.topLevelItem(i)
            config_dict["Key_parameters_single"].append([])
            for k in range(item.columnCount()):
                config_dict["Key_parameters_single"][-1].append(item.text(k))
        config_dict["Key_parameters_multy"] = []
        for i in range(self.ui.treeWidget_key_parameters_multy.topLevelItemCount()):
            item = self.ui.treeWidget_key_parameters_multy.topLevelItem(i)
            config_dict["Key_parameters_multy"].append([])
            for k in range(item.columnCount()):
                config_dict["Key_parameters_multy"][-1].append(item.text(k))


        file_name = QFileDialog.getSaveFileName(self, "Сохранить файл", "", "Файлы конфигурации (*.json)")[0]
        if len(file_name) > 0:
            with open(file_name, "w", encoding="utf8") as config_file:
                json.dump(config_dict, config_file, indent=4, ensure_ascii=False)

    def open_sfg(self):
        path = QFileDialog.getOpenFileName(self, "Открыть файл", "", "Файлы конфигурации (*.json)")[0]
        if len(path) > 0:
            with open(path, encoding="utf8") as json_file:
                config_dict=json.load(json_file)

            # Корректируемые параметры
            self.ui.treeWidget_corrected_parameters.clear()
            for cfg_columns in config_dict["Correcting"]["Items"]:
                item = QTreeWidgetItem(cfg_columns)
                self.ui.treeWidget_corrected_parameters.addTopLevelItem(item)
            # Нужно для completer
            self.ui.treeWidget_corrected_parameters.setCurrentItem(item)

            # Графики
            self.ui.treeWidget_graphs_single.clear()
            for cfg_columns in config_dict["Graphs_single"]:
                item = QTreeWidgetItem(cfg_columns)
                self.ui.treeWidget_graphs_single.addTopLevelItem(item)

            self.ui.treeWidget_graphs_multy.clear()
            for cfg_columns in config_dict["Graphs_multy"]:
                item = QTreeWidgetItem(cfg_columns)
                self.ui.treeWidget_graphs_multy.addTopLevelItem(item)

            # Хронологии
            self.ui.treeWidget_chrono_single.clear()
            for cfg_columns in config_dict["Chrono_single"]:
                item = QTreeWidgetItem(cfg_columns)
                self.ui.treeWidget_chrono_single.addTopLevelItem(item)

            self.ui.treeWidget_chrono_multy.clear()
            for cfg_columns in config_dict["Chrono_multy"]:
                item = QTreeWidgetItem(cfg_columns)
                self.ui.treeWidget_chrono_multy.addTopLevelItem(item)

            # Ключевые параметры
            self.ui.treeWidget_key_parameters_single.clear()
            for cfg_columns in config_dict["Key_parameters_single"]:
                item = QTreeWidgetItem(cfg_columns)
                self.ui.treeWidget_key_parameters_single.addTopLevelItem(item)

            self.ui.treeWidget_key_parameters_multy.clear()
            for cfg_columns in config_dict["Key_parameters_multy"]:
                item = QTreeWidgetItem(cfg_columns)
                self.ui.treeWidget_key_parameters_multy.addTopLevelItem(item)

    # метод для редактирования корректируемого параметра
    def edit_correcting(self, item, column):
        dlg = CorrectingEditDlg(item)
        self.completer.update_correcting_parameters(self.ui.treeWidget_corrected_parameters)
        dlg.line_corrected_parameters_name.setCompleter(self.completer.clusters_completer)
        dlg.exec()
    # Вкладка Графики
    # Метод для выбора дерева графиков
    def get_tree_graphs(self):
        tree = None
        if self.ui.tab_graphs_list.currentIndex() == 0:
            tree = self.ui.treeWidget_graphs_single
        elif self.ui.tab_graphs_list.currentIndex() == 1:
            tree = self.ui.treeWidget_graphs_multy
        if tree is not None:
            return tree

    # метод для сброса индекса графиков
    def reset_index_graphs(self):
        tree = self.get_tree_graphs()
        if tree is not None:
            items = tree.findItems("", Qt.MatchFlag.MatchContains)
            for i in range(len(items)):
                items[i].setText(0, str(i + 1))

    # метод для добавления графиков
    def add_graph(self):
        tree = self.get_tree_graphs()
        if tree is not None:
            number = str(len(tree.findItems("", Qt.MatchFlag.MatchContains)) + 1)
            name_y = self.ui.line_graphs_y_names.text()
            name_x = self.ui.line_graphs_x_names.text()
            label_y = self.ui.line_graphs_y_label.text()
            label_x = self.ui.line_graphs_x_label.text()
            min_y = self.ui.line_graphs_y_min.text()
            max_y = self.ui.line_graphs_y_max.text()
            min_x = self.ui.line_graphs_x_min.text()
            max_x = self.ui.line_graphs_x_max.text()
            mult_y = self.ui.line_graphs_y_mult.text()
            mult_x = self.ui.line_graphs_x_mult.text()
            shift_y = self.ui.line_graphs_y_shift.text()
            shift_x = self.ui.line_graphs_x_shift.text()
            step_y = self.ui.line_graphs_y_step.text()
            step_x = self.ui.line_graphs_x_step.text()
            description = self.ui.text_edit_graphs_description.toPlainText()
            item = QTreeWidgetItem([str(number),
                                    name_y,
                                    name_x,
                                    label_y,
                                    label_x,
                                    min_y,
                                    max_y,
                                    min_x,
                                    max_x,
                                    mult_y,
                                    shift_y,
                                    mult_x,
                                    shift_x,
                                    step_y,
                                    step_x,
                                    description,
                                    ])
            tree.addTopLevelItem(item)

    # метод для удаления выбранных графиков
    def del_selected_graphs(self):
        tree = self.get_tree_graphs()
        if tree is not None:
            self.del_selected_elements(tree)
            self.reset_index_graphs()

    # метод для перемещения графика вверх
    def move_up_selected_graphs(self):
        tree = self.get_tree_graphs()
        if tree is not None:
            self.move_selected_up(tree)
            self.reset_index_graphs()

    # метод для перемещения графика вниз
    def move_down_selected_graphs(self):
        tree = self.get_tree_graphs()
        if tree is not None:
            self.move_selected_down(tree)
            self.reset_index_graphs()
    # метод для редактирования графика
    def edit_graph(self, item, column):
        dlg=GraphEditDlg(item)
        # Графики
        dlg.line_graphs_x_names.setCompleter(self.completer.graphs_completer)
        dlg.line_graphs_x_names.textChanged.connect(
            lambda: self.completer.update_graphs(dlg.line_graphs_y_names))

        dlg.line_graphs_y_names.setCompleter(self.completer.graphs_completer)
        dlg.line_graphs_y_names.textChanged.connect(
            lambda: self.completer.update_graphs(dlg.line_graphs_y_names))
        dlg.exec()

    # Вкладка Хронология
    # Метод для выбора дерева хронологий
    def get_tree_chrono(self):
        tree = None
        if self.ui.tab_chrono_list.currentIndex() == 0:
            tree = self.ui.treeWidget_chrono_single
        elif self.ui.tab_chrono_list.currentIndex() == 1:
            tree = self.ui.treeWidget_chrono_multy
        if tree is not None:
            return tree

    # Метод для добавления события
    def add_chrono(self):
        tree = self.get_tree_chrono()
        if tree is not None:
            name = self.ui.line_chrono_parameter_name.text()
            value = self.ui.line_chrono_parameter_value.text()
            cut_down = ";".join([self.ui.line_cut_down_chrono_name.text(), self.ui.line_cut_down_chrono_value.text()])
            cut_up = ";".join([self.ui.line_cut_up_chrono_name.text(), self.ui.line_cut_up_chrono_value.text()])
            description = self.ui.text_edit_chrono_description.toPlainText()
            if name == "":
                print("Введено пустое название")
                self.show_message_window("Введено пустое название")
            elif value == "":
                print("Введено пустое значение")
                self.show_message_window("Введено пустое значение")
            elif description == "":
                print("Введено пустое описание")
                self.show_message_window("Введено пустое описание")
            else:
                item = QTreeWidgetItem([name,
                                        value,
                                        cut_down,
                                        cut_up,
                                        description,
                                        ])
                tree.addTopLevelItem(item)
                self.ui.line_chrono_parameter_name.setText("")
                self.ui.line_chrono_parameter_value.setText("")
                self.ui.text_edit_chrono_description.setText("")
                self.ui.line_cut_down_chrono_name.setText("")
                self.ui.line_cut_down_chrono_value.setText("")
                self.ui.line_cut_up_chrono_name.setText("")
                self.ui.line_cut_up_chrono_value.setText("")

    # метод для редактирования хронологии
    def edit_chrono(self, item, column):
        dlg=ChronoEditDlg(item)

        # Хронология
        dlg.line_chrono_parameter_name.textChanged.connect(lambda:
                                                               self.selection_completer_clusters_report(
                                                                   dlg.line_chrono_parameter_name,
                                                                   dlg.line_chrono_parameter_value))
        dlg.line_chrono_parameter_value.textChanged.connect(lambda:
                                                                self.selection_completer_help(
                                                                    dlg.line_chrono_parameter_name,
                                                                    dlg.line_chrono_parameter_value))
        # Обрезка хронологий снизу
        dlg.line_cut_down_chrono_name.setCompleter(self.completer.clusters_completer)
        dlg.line_cut_down_chrono_value.setCompleter(self.completer.help_clusters_completer)
        # Обрезка хронологий сверху
        dlg.line_cut_up_chrono_name.setCompleter(self.completer.clusters_completer)
        dlg.line_cut_up_chrono_value.setCompleter(self.completer.help_clusters_completer)

        dlg.exec()

    # Вкладка Ключевые параметры
    # Метод для выбора дерева ключевых параметров
    def get_tree_key_parameters(self):
        tree = None
        if self.ui.tab_key_parameters_list.currentIndex() == 0:
            tree = self.ui.treeWidget_key_parameters_single
        elif self.ui.tab_key_parameters_list.currentIndex() == 1:
            tree = self.ui.treeWidget_key_parameters_multy
        if tree is not None:
            return tree

    # Метод для добавления ключевых параметров
    def add_key_parameter(self):
        tree = self.get_tree_key_parameters()
        if tree is not None:
            key_parameter_name = self.ui.line_key_parameter_name.text()
            search_parameter_name = self.ui.line_search_parameter_name.text()
            value = self.ui.line_search_parameter_value.text()
            cut_down = ";".join(
                [self.ui.line_cut_down_key_parameters_name.text(), self.ui.line_cut_down_key_parameters_value.text()])
            cut_up = ";".join(
                [self.ui.line_cut_up_key_parameters_name.text(), self.ui.line_cut_up_key_parameters_value.text()])
            description = self.ui.text_edit_key_parameters_description.toPlainText()
            if key_parameter_name == "":
                print("Введено пустое название ключевого параметра")
                self.show_message_window("Введено пустое название ключевого параметра")
            elif search_parameter_name == "":
                print("Введено пустое название параметра для поиска")
                self.show_message_window("Введено пустое название параметра для поиска")
            elif value == "":
                print("Введено пустое значение")
                self.show_message_window("Введено пустое значение")
            elif description == "":
                print("Введено пустое описание")
                self.show_message_window("Введено пустое описание")
            else:
                item = QTreeWidgetItem([key_parameter_name,
                                        search_parameter_name,
                                        value,
                                        cut_down,
                                        cut_up,
                                        description,
                                        ])
                tree.addTopLevelItem(item)
                self.ui.line_key_parameter_name.setText("")
                self.ui.line_search_parameter_name.setText("")
                self.ui.line_search_parameter_value.setText("")
                self.ui.text_edit_key_parameters_description.setText("")
                self.ui.line_cut_down_key_parameters_name.setText("")
                self.ui.line_cut_down_key_parameters_value.setText("")
                self.ui.line_cut_up_key_parameters_name.setText("")
                self.ui.line_cut_up_key_parameters_value.setText("")

    # метод для перемещения ключевого параметра вверх
    def move_up_selected_key_parameters(self):
        tree = self.get_tree_key_parameters()
        if tree is not None:
            self.move_selected_up(tree)

    # метод для перемещения ключевого параметра вниз
    def move_down_selected_key_parameters(self):
        tree = self.get_tree_key_parameters()
        if tree is not None:
            self.move_selected_down(tree)

    # метод для редактирования ключевого параметра
    def edit_key_parameter(self, item, column):
        dlg=KeyParametersEditDlg(item)

        dlg.line_key_parameter_name.setCompleter(self.completer.clusters_completer)
        dlg.line_search_parameter_name.textChanged.connect(lambda:
                                                               self.selection_completer_clusters_report(
                                                                   dlg.line_search_parameter_name,
                                                                   dlg.line_search_parameter_value))
        dlg.line_search_parameter_value.textChanged.connect(lambda:
                                                                self.selection_completer_help(
                                                                    dlg.line_search_parameter_name,
                                                                    dlg.line_search_parameter_value))
        # Обрезка снизу ключевых параметров
        dlg.line_cut_down_key_parameters_name.setCompleter(self.completer.clusters_completer)
        dlg.line_cut_down_key_parameters_value.setCompleter(self.completer.help_clusters_completer)
        # Обрезка сверху ключевых параметров
        dlg.line_cut_up_key_parameters_name.setCompleter(self.completer.clusters_completer)
        dlg.line_cut_up_key_parameters_value.setCompleter(self.completer.help_clusters_completer)

        dlg.exec()

    # метод для предварительной загрузки
    def pre_run(self):
        self.stopped = False
        # Список для хранения прочитанных данных
        dict_list=[]
        # функция для обработки cut_down и cut_up
        def cut_str_to_dict(cut_str):
            name, value = cut_str.split(";")
            if value != "min" and value != "max":
                try:
                    value = float(value)
                except:
                    value = None

            return {"name": name, "value": value}


        # Получение параметров с вкладки Данные
        calculations_columns = ["name", "cut_down", "cut_up", "path"]
        for item_index in range(self.ui.treeWidget_calculations.topLevelItemCount()):
            item = self.ui.treeWidget_calculations.topLevelItem(item_index)
            item_dict = {}
            item_values = [item.text(columns_index) for columns_index in range(item.columnCount())]
            for column, value in zip(calculations_columns, item_values):
                item_dict[column] = value
            for cut_name in ("cut_down", "cut_up"):
                item_dict[cut_name] = cut_str_to_dict(item_dict[cut_name])

            dict_list.append(item_dict)

        # Список с прочитаными расчётами. Удаляет "calc_object" из self.calculations
        calc_object_list=[calculation.pop("calc_object") for calculation in self.calculations]

        for i in range(len(self.calculations)):
            if self.calculations[i] in dict_list:
                k=dict_list.index(self.calculations[i])
                dict_list[k]["calc_object"] = calc_object_list[i]

        for item_dict in dict_list:
            if self.stopped:
                return
            if not "calc_object" in item_dict:
                if item_dict["path"].endswith(".sokle"):
                    with open(item_dict["path"], 'rb') as f:
                        calc = pickle.load(f)
                elif os.path.isdir(item_dict["path"]):
                    file_names_list = os.listdir(item_dict["path"])
                    file_names = ";".join(file_names_list)
                    if ".dia;" in file_names:
                        calc = Socrat_calculation(item_dict["path"])
                    elif "lent3" in file_names_list:
                        calc = Trap_calculation(item_dict["path"])
                    else:
                        print("Проблемы с " + item_dict["path"])
                        self.show_message_window("Проблемы с " + item_dict["path"])
                        self.calculations = []
                        break
                else:
                    print("Проблемы с " + item_dict["path"])
                    self.show_message_window("Проблемы с " + item_dict["path"])
                    self.calculations = []
                    break

                # Заполнение множеств для автодополнения
                for table_name in calc.data_dict.keys():
                    if table_name == "report":
                        self.completer.update_report(calc.data_dict[table_name]["Name_report"])
                    else:
                        self.completer.update_clusters(calc.data_dict[table_name].columns)

                item_dict["calc_object"]=calc

        self.calculations = dict_list
    # метод для БОЛЬШОЙ КНОПКИ Выполнить
    def run_programm(self):
        self.stopped = False
        # функция для обработки cut_down и cut_up
        def cut_str_to_dict(cut_str):
            name, value = cut_str.split(";")
            if value != "min" and value != "max":
                try:
                    value = float(value)
                except:
                    value = None

            return {"name": name, "value": value}

        # Список для Multy
        calculations_list = []

        path_save = self.ui.line_path_to_save.text()

        # функция для перевода из строки в float число
        def str_to_number(number_str):
            try:
                return float(number_str)
            except:
                return None

        # функция для обработки строки со значением для поиска
        def search_str_transform(search_value):
            if not ((("ON" in search_value or "OFF" in search_value) and (
                    "FIRST" in search_value or "LAST" in search_value)) or search_value == "min" or search_value == "max"):
                try:
                    search_value = float(search_value)
                except:
                    search_value = None
            return search_value

        # функции для округления ключевых и хронологических чисел
        def round_chrono(value):
            abs_value = abs(value)
            if abs_value < 10:
                return round(value, 2)
            elif abs_value < 100:
                return round(value, 1)
            else:
                return int(round(value, 0))
        def round_key(value):
            abs_value = abs(value)
            if abs_value < 10:
                return round(value, 2)
            elif abs_value < 100:
                return round(value, 1)
            else:
                return int(round(value, 0))

        # Предварительная загрузка:
        self.pre_run()

        # Получение параметров с вкладки Обработка
        corrected_parameters = []
        corrected_parameters_columns = ["name", "extra_operations", "expression"]
        for item_index in range(self.ui.treeWidget_corrected_parameters.topLevelItemCount()):
            item = self.ui.treeWidget_corrected_parameters.topLevelItem(item_index)
            item_dict = {}
            item_values = [item.text(columns_index) for columns_index in range(item.columnCount())]
            for column, value in zip(corrected_parameters_columns, item_values):
                item_dict[column] = value

            corrected_parameters.append(item_dict)

        # Получение параметров с вкладки Графики
        graphs_single = []
        graphs_multy = []
        graphs_columns = ["number",
                          "y_name", "x_name",
                          "y_label", "x_label",
                          "y_cut_down", "y_cut_up",
                          "x_cut_down", "x_cut_up",
                          "y_mult", "y_shift",
                          "x_mult", "x_shift",
                          "y_step", "x_step",
                          "description"]

        for item_index in range(self.ui.treeWidget_graphs_single.topLevelItemCount()):
            item = self.ui.treeWidget_graphs_single.topLevelItem(item_index)
            item_dict = {}
            item_values = [item.text(columns_index) for columns_index in range(item.columnCount())]
            for column, value in zip(graphs_columns, item_values):
                item_dict[column] = value
            for split_name in ("y_name", "x_name"):
                item_dict[split_name] = item_dict[split_name].split(";")
            for number_name in (
                    "y_cut_down", "y_cut_up", "x_cut_down", "x_cut_up", "y_mult", "y_shift", "x_mult", "x_shift",
                    "y_step", "x_step",):
                item_dict[number_name] = str_to_number(item_dict[number_name])

            graphs_single.append(item_dict)

        for item_index in range(self.ui.treeWidget_graphs_multy.topLevelItemCount()):
            item = self.ui.treeWidget_graphs_multy.topLevelItem(item_index)
            item_dict = {}
            item_values = [item.text(columns_index) for columns_index in range(item.columnCount())]
            for column, value in zip(graphs_columns, item_values):
                item_dict[column] = value
            for split_name in ("y_name", "x_name"):
                item_dict[split_name] = item_dict[split_name].split(";")
            for number_name in (
                    "y_cut_down", "y_cut_up", "x_cut_down", "x_cut_up", "y_mult", "y_shift", "x_mult", "x_shift",
                    "y_step",
                    "x_step"):
                item_dict[number_name] = str_to_number(item_dict[number_name])
            graphs_multy.append(item_dict)

        # Получение параметров с вкладки Хронология
        chrono_single = []
        chrono_multy = []
        chrono_columns = ["name", "value", "cut_down", "cut_up", "description"]

        for item_index in range(self.ui.treeWidget_chrono_single.topLevelItemCount()):
            item = self.ui.treeWidget_chrono_single.topLevelItem(item_index)
            item_dict = {}
            item_values = [item.text(columns_index) for columns_index in range(item.columnCount())]
            for column, value in zip(chrono_columns, item_values):
                item_dict[column] = value
            item_dict["value"] = search_str_transform(item_dict["value"])
            for cut_name in ("cut_down", "cut_up"):
                item_dict[cut_name] = cut_str_to_dict(item_dict[cut_name])
            chrono_single.append(item_dict)

        for item_index in range(self.ui.treeWidget_chrono_multy.topLevelItemCount()):
            item = self.ui.treeWidget_chrono_multy.topLevelItem(item_index)
            item_dict = {}
            item_values = [item.text(columns_index) for columns_index in range(item.columnCount())]
            for column, value in zip(chrono_columns, item_values):
                item_dict[column] = value
            item_dict["value"] = search_str_transform(item_dict["value"])
            for cut_name in ("cut_down", "cut_up"):
                item_dict[cut_name] = cut_str_to_dict(item_dict[cut_name])
            chrono_multy.append(item_dict)

        # Получение параметров с вкладки Ключевые параметры
        key_parameters_single = []
        key_parameters_multy = []
        key_parameters_columns = ["key_parameter_name", "search_name", "value", "cut_down", "cut_up", "description"]

        for item_index in range(self.ui.treeWidget_key_parameters_single.topLevelItemCount()):
            item = self.ui.treeWidget_key_parameters_single.topLevelItem(item_index)
            item_dict = {}
            item_values = [item.text(columns_index) for columns_index in range(item.columnCount())]
            for column, value in zip(key_parameters_columns, item_values):
                item_dict[column] = value
            item_dict["value"] = search_str_transform(item_dict["value"])
            for cut_name in ("cut_down", "cut_up"):
                item_dict[cut_name] = cut_str_to_dict(item_dict[cut_name])
            key_parameters_single.append(item_dict)

        for item_index in range(self.ui.treeWidget_key_parameters_multy.topLevelItemCount()):
            item = self.ui.treeWidget_key_parameters_multy.topLevelItem(item_index)
            item_dict = {}
            item_values = [item.text(columns_index) for columns_index in range(item.columnCount())]
            for column, value in zip(key_parameters_columns, item_values):
                item_dict[column] = value
            item_dict["value"] = search_str_transform(item_dict["value"])
            for cut_name in ("cut_down", "cut_up"):
                item_dict[cut_name] = cut_str_to_dict(item_dict[cut_name])
            key_parameters_multy.append(item_dict)

        # Получение параметров с вкладки Сохранение результатов
        file_save_name = self.ui.line_file_save_name.text()
        if file_save_name == "":
            file_save_name = "Results"
        graphs_width = self.ui.line_graphs_width.text()
        if graphs_width == "":
            graphs_width = 12
        else:
            graphs_width = float(graphs_width)
        number_delim = self.ui.line_numbers_delim.text()
        font_name = self.ui.comboBox_font_name.currentText()
        if not font_name:
            font_name = "Times New Roman"
        font_value = self.ui.comboBox_font_value.currentText()
        if font_value.isdigit():
            font_value = float(font_value)
        else:
            font_value = 12
        check_round_chrono = self.ui.checkBox_round_chrono.checkState()
        check_round_key_parameters = self.ui.checkBox_round_key_parameters.checkState()

        # Открытие папки для сохранения
        if len(path_save) == 0:
            os.chdir("..")
            if not os.path.isdir("Socrat_eye_output"):
                os.mkdir("Socrat_eye_output")
            path_save = os.path.abspath("Socrat_eye_output")
            self.ui.line_path_to_save.setText(path_save)
        os.chdir(path_save)
        # Обработка результатов
        for calculation in self.calculations:
            if self.stopped:
                return

            calc=calculation["calc_object"]
            if self.stopped:
                return
            # Обрезка снизу и сверху
            calc.cut_bottom_transform(calculation["cut_down"]["name"], calculation["cut_down"]["value"])
            calc.cut_top_transform(calculation["cut_up"]["name"], calculation["cut_up"]["value"])
            if self.stopped:
                return
            # При необходимости сохранение в бинарном виде
            if not calculation["path"].endswith(".sokle"):
                calc.save(name=calculation["name"])

            # Корректировка
            for corrected_parameter in corrected_parameters:
                if self.stopped:
                    return
                integral = False
                derivative = False
                if corrected_parameter["extra_operations"]==self.ui.radioButton_integral.text():
                    integral=True
                elif corrected_parameter["extra_operations"]==self.ui.radioButton_gradient.text():
                    derivative = True

                calc.correcting_parameters(corrected_parameter["name"], corrected_parameter["expression"],
                                       integral=integral, derivative=derivative)


            # Построение графиков
            if len(graphs_single)>0:
                # Открытие папки для сохранения в режиме Single
                if not os.path.exists(calculation["name"]):
                    os.mkdir(calculation["name"])
                os.chdir(calculation["name"])
                # Построение графиков Single
                for graph_single in graphs_single:
                    if self.stopped:
                        return
                    print(graph_single)
                    calc.graph(GrName=graph_single["number"],
                               x_names=graph_single["x_name"],
                               y_names=graph_single["y_name"],
                               lablex=graph_single["x_label"],
                               labley=graph_single["y_label"],
                               x1=graph_single["x_cut_down"],
                               x2=graph_single["x_cut_up"],
                               y1=graph_single["y_cut_down"],
                               y2=graph_single["y_cut_up"],
                               mult_x=graph_single["x_mult"],
                               mult_y=graph_single["y_mult"],
                               shift_x=graph_single["x_shift"],
                               shift_y=graph_single["y_shift"],
                               stpx=graph_single["x_step"],
                               stpy=graph_single["y_step"],
                               empty_graphs=self.ui.checkBox_empty_graphs.isChecked())

            if self.stopped:
                return
            # Построение хронологий и ключевых параметров Single
            calc.make_chrono(chrono_single)
            if self.stopped:
                return
            calc.make_key_table(key_parameters_single)

            # Округление Single
            if check_round_chrono:
                calc.chrono_table["Time"] = calc.chrono_table["Time"].apply(round_chrono)
            if check_round_key_parameters:
                calc.key_parameters_table["Value"] = calc.key_parameters_table["Value"].apply(round_key)

            # Возврат в корневую папку
            os.chdir(path_save)

            # Сохранение в Word Single
            doc = docx.Document()
            style = doc.styles['Normal']
            style.font.name = font_name
            style.font.size = docx.shared.Pt(font_value)

            # Сохранение хронологии Single
            if len(calc.chrono_table) > 0:
                doc.add_paragraph("Хронология ")
                table = doc.add_table(rows=1, cols=2)
                table.style = "Table Grid"
                table.cell(0, 0).text = "Время, с"
                table.cell(0, 1).text = "Событие"
                for k in range(len(calc.chrono_table)):
                    table.add_row()
                    value_str=str(calc.chrono_table.iloc[k, 0])
                    if value_str.endswith(".0"):
                        value_str=value_str[:-2]
                    if number_delim!=".":
                        value_str=value_str.replace(".", number_delim)
                    table.cell(k + 1, 0).text = value_str
                    table.cell(k + 1, 1).text = str(calc.chrono_table.iloc[k, 1])
                doc.add_page_break()

            # Сохранение ключевых параметров Single
            if len(calc.key_parameters_table) > 0:
                doc.add_paragraph("Ключевые параметры")
                table = doc.add_table(rows=1, cols=2)
                table.style = "Table Grid"
                table.cell(0, 0).text = "Параметр"
                table.cell(0, 1).text = "Значение"
                for k in range(len(calc.key_parameters_table)):
                    table.add_row()
                    table.cell(k + 1, 0).text = str(calc.key_parameters_table.iloc[k, 0])
                    table.cell(k + 1, 1).text = str(calc.key_parameters_table.iloc[k, 1]).replace(".", number_delim)
                doc.add_page_break()

            # Сохранение графиков Single
            for k in range(len(calc.graph_list)):
                path_to_png = calc.graph_list[k]["png"]
                index_png = int(os.path.basename(path_to_png).replace(".png",""))-1
                doc.add_picture(path_to_png, width=docx.shared.Cm(graphs_width))
                doc.add_paragraph(
                    "Рисунок {}\n{}".format(str(k + 1), graphs_single[index_png]["description"]))

            # Запись Word Single
            if len(calc.graph_list)>0 or len(calc.key_parameters_table) or len(calc.chrono_table) >0:
                doc_saved=False
                doc_name=calculation["name"] + ".docx"
                while not doc_saved:
                    try:
                        doc.save(doc_name)
                        doc_saved=True
                    except:
                        print("Файл {} в данный момент открыт".format(doc_name))
                        self.show_message_window("Файл {} в данный момент открыт".format(doc_name))

            # Сброс графиков, хронологий и ключевых параметров
            calc.reset_data()

            # Построение хронологий и ключевых параметров Multy
            if self.stopped:
                return
            calc.make_chrono(chrono_multy)
            if self.stopped:
                return
            calc.make_key_table(key_parameters_multy)

            # Округление Multy
            if check_round_chrono:
                calc.chrono_table["Time"] = calc.chrono_table["Time"].apply(round_chrono)
            if check_round_key_parameters:
                calc.key_parameters_table["Value"] = calc.key_parameters_table["Value"].apply(round_key)

            calculations_list.append(calc)

        # Создание объекта для хранения Multy
        calculation_series = Series_of_calculations(calculations_list)

        # Построение графиков
        if len(graphs_multy)>0:
            # Открытие папки для сохранения в режиме Multy
            if not os.path.exists(file_save_name):
                os.mkdir(file_save_name)
            os.chdir(file_save_name)
            # Отрисовка графиков Multy
            for graph_multy in graphs_multy:
                if self.stopped:
                    return
                calculation_series.graph(GrName=graph_multy["number"],
                                         x_names=graph_multy["x_name"],
                                         y_names=graph_multy["y_name"],
                                         lablex=graph_multy["x_label"],
                                         labley=graph_multy["y_label"],
                                         x1=graph_multy["x_cut_down"],
                                         x2=graph_multy["x_cut_up"],
                                         y1=graph_multy["y_cut_down"],
                                         y2=graph_multy["y_cut_up"],
                                         mult_x=graph_multy["x_mult"],
                                         mult_y=graph_multy["y_mult"],
                                         shift_x=graph_multy["x_shift"],
                                         shift_y=graph_multy["y_shift"],
                                         stpx=graph_multy["x_step"],
                                         stpy=graph_multy["y_step"],
                                         empty_graphs=self.ui.checkBox_empty_graphs.isChecked())

        # Возврат в корневую папку
        os.chdir(path_save)

        # Сохранение в Word Multy
        doc = docx.Document()
        style = doc.styles['Normal']
        style.font.name = font_name
        style.font.size = docx.shared.Pt(font_value)

        # Расшифровка чисел
        numbers_to_names=["{} - {}".format(str(i+1),self.calculations[i]["name"]) for i in range(len(self.calculations))]
        numbers_to_names="\n".join(numbers_to_names)
        doc.add_paragraph(numbers_to_names)


        # Сохранение хронологий Multy
        if len(calculation_series.chrono_table) > 0:
            doc.add_page_break()
            doc.add_paragraph("Хронология")
            table = doc.add_table(rows=len(calculation_series.chrono_table) + 1,
                                  cols=len(calculation_series.chrono_table.columns))
            table.style = "Table Grid"
            table.cell(0, 0).text = "Событие"
            for i in range(len(self.calculations)):
                table.cell(0, i + 1).text = str(i+1)
            for i in range(len(calculation_series.chrono_table.columns)):
                for k in range(len(calculation_series.chrono_table)):
                    table.cell(k + 1, i).text = str(calculation_series.chrono_table.iloc[k, i]).replace(".", number_delim)


        # Сохранение ключевых параметров Multy
        if len(calculation_series.key_parameters_table) > 0:
            doc.add_page_break()
            doc.add_paragraph("Ключевые параметры")
            table = doc.add_table(rows=len(calculation_series.key_parameters_table) + 1,
                                  cols=len(calculation_series.key_parameters_table.columns))
            table.style = "Table Grid"
            table.cell(0, 0).text = "Параметр"
            for i in range(len(self.calculations)):
                table.cell(0, i + 1).text = str(i+1)
            for i in range(len(calculation_series.key_parameters_table.columns)):
                for k in range(len(calculation_series.key_parameters_table)):
                    table.cell(k + 1, i).text = str(calculation_series.key_parameters_table.iloc[k, i]).replace(".", number_delim)

        # Сохранение графиков Multy
        if len(calculation_series.graph_list) > 0:
            doc.add_page_break()
            for k in range(len(calculation_series.graph_list)):
                path_to_png = calculation_series.graph_list[k]["png"]
                index_png = int(os.path.basename(path_to_png).replace(".png", ""))-1
                doc.add_picture(path_to_png, width=docx.shared.Cm(graphs_width))
                doc.add_paragraph("Рисунок {}\n{}".format(str(k + 1), graphs_multy[index_png]["description"]))

        # Запись Word Multy
        if len(calculation_series.graph_list) > 0 or len(calculation_series.key_parameters_table) or len(calculation_series.chrono_table) > 0:
            doc_saved = False
            doc_name = file_save_name + ".docx"
            while not doc_saved:
                try:
                    doc.save(file_save_name + ".docx")
                    doc_saved = True
                except:
                    print("Файл {} в данный момент открыт".format(doc_name))
                    self.show_message_window("Файл {} в данный момент открыт".format(doc_name))
        os.startfile(path_save)
    # метод для запуска выполнения на отдельном потоке выполнения программы
    def run_function_safely(self, fn):
        worker = Worker(fn)
        self.set_interface_freezed(True)
        # worker.signals.result.connect(lambda: self.show_message_window("Готово"))
        worker.signals.finished.connect(lambda: self.set_interface_freezed(False))
        self.thread_manager.start(worker)
    # метод для заморозки интерфейса
    def set_interface_freezed(self, state):
        if state:
            self.ui.button_execute.setDisabled(True)
            self.ui.button_pre_run.setDisabled(True)
            self.ui.button_path_save.setDisabled(True)
            self.ui.line_path_to_save.setDisabled(True)
            for i in range(self.ui.tabWidget_main.count()):
                self.ui.tabWidget_main.widget(i).setDisabled(True)
        else:
            self.ui.button_execute.setEnabled(True)
            self.ui.button_pre_run.setEnabled(True)
            self.ui.button_path_save.setEnabled(True)
            self.ui.line_path_to_save.setEnabled(True)
            for i in range(self.ui.tabWidget_main.count()):
                self.ui.tabWidget_main.widget(i).setEnabled(True)
    # метод для остановки расчета
    def stop_run(self):
        self.stopped = True

# Обертка для перенаправления потока вывода
class OutputWrapper(QObject):
    outputWritten = Signal(object,object)

    def __init__(self, parent, stdout = True):
        super().__init__(parent)
        if stdout:
            self._stream = sys.stdout
            sys.stdout = self
        else:
            self._stream = sys.stderr
            sys.stderr = self
        self._stdout = stdout
    def write(self, text):
        self._stream.write(text)
        self.outputWritten.emit(text, self._stdout)

    def __getattr__(self, name):
        return getattr(self._stream, name)

    def __del__(self):
        try:
            if self._stdout:
                sys.stdout = self._stream
            else:
                sys.stderr = self._stream
        except AttributeError:
            pass


class GraphEditDlg(Ui_Dialog_graphs_edit, QDialog):
    def __init__(self, item, parent=None):
        super().__init__(parent)
        self.setupUi(self)
        self.item=item
        self.line_graphs_y_names.setText(item.text(1))
        self.line_graphs_x_names.setText(item.text(2))
        self.line_graphs_y_label.setText(item.text(3))
        self.line_graphs_x_label.setText(item.text(4))
        self.line_graphs_y_min.setText(item.text(5))
        self.line_graphs_y_max.setText(item.text(6))
        self.line_graphs_x_min.setText(item.text(7))
        self.line_graphs_x_max.setText(item.text(8))
        self.line_graphs_y_mult.setText(item.text(9))
        self.line_graphs_y_shift.setText(item.text(10))
        self.line_graphs_x_mult.setText(item.text(11))
        self.line_graphs_x_shift.setText(item.text(12))
        self.line_graphs_y_step.setText(item.text(13))
        self.line_graphs_x_step.setText(item.text(14))
        self.text_edit_graphs_description.setText(item.text(15))

        self.buttonBox.accepted.connect(self.save_changes)

    def save_changes(self):
        self.item.setText(1, self.line_graphs_y_names.text())
        self.item.setText(2, self.line_graphs_x_names.text())
        self.item.setText(3, self.line_graphs_y_label.text())
        self.item.setText(4, self.line_graphs_x_label.text())
        self.item.setText(5, self.line_graphs_y_min.text())
        self.item.setText(6, self.line_graphs_y_max.text())
        self.item.setText(7, self.line_graphs_x_min.text())
        self.item.setText(8, self.line_graphs_x_max.text())
        self.item.setText(9, self.line_graphs_y_mult.text())
        self.item.setText(10, self.line_graphs_y_shift.text())
        self.item.setText(11, self.line_graphs_x_mult.text())
        self.item.setText(12, self.line_graphs_x_shift.text())
        self.item.setText(13, self.line_graphs_y_step.text())
        self.item.setText(14, self.line_graphs_x_step.text())
        self.item.setText(15, self.text_edit_graphs_description.toPlainText())
class ChronoEditDlg(Ui_Dialog_chrono_edit, QDialog):
    def __init__(self, item, parent=None):
        super().__init__(parent)
        self.setupUi(self)
        self.item=item


        self.line_chrono_parameter_name.setText(item.text(0))
        self.line_chrono_parameter_value.setText(item.text(1))
        self.line_cut_down_chrono_name.setText(item.text(2).split(";")[0])
        self.line_cut_down_chrono_value.setText(item.text(2).split(";")[1])
        self.line_cut_up_chrono_name.setText(item.text(3).split(";")[0])
        self.line_cut_up_chrono_value.setText(item.text(3).split(";")[1])

        self.text_edit_chrono_description.setText(item.text(4))

        self.buttonBox.accepted.connect(self.save_changes)

    def save_changes(self):
        self.item.setText(0, self.line_chrono_parameter_name.text())
        self.item.setText(1, self.line_chrono_parameter_value.text())
        self.item.setText(2, ";".join([self.line_cut_down_chrono_name.text(),self.line_cut_down_chrono_value.text()]))
        self.item.setText(3, ";".join([self.line_cut_up_chrono_name.text(), self.line_cut_up_chrono_value.text()]))
        self.item.setText(4, self.text_edit_chrono_description.toPlainText())
class KeyParametersEditDlg(Ui_Dialog_key_parameters_edit, QDialog):
    def __init__(self, item, parent=None):
        super().__init__(parent)
        self.setupUi(self)
        self.item=item

        self.line_key_parameter_name.setText(item.text(0))
        self.line_search_parameter_name.setText(item.text(1))
        self.line_search_parameter_value.setText(item.text(2))
        self.line_cut_down_key_parameters_name.setText(item.text(3).split(";")[0])
        self.line_cut_down_key_parameters_value.setText(item.text(3).split(";")[1])
        self.line_cut_up_key_parameters_name.setText(item.text(4).split(";")[0])
        self.line_cut_up_key_parameters_value.setText(item.text(4).split(";")[1])

        self.text_edit_key_parameters_description.setText(item.text(5))

        self.buttonBox.accepted.connect(self.save_changes)

    def save_changes(self):
        self.item.setText(0, self.line_key_parameter_name.text())
        self.item.setText(1, self.line_search_parameter_name.text())
        self.item.setText(2, self.line_search_parameter_value.text())
        self.item.setText(3, ";".join([self.line_cut_down_key_parameters_name.text(),self.line_cut_down_key_parameters_value.text()]))
        self.item.setText(4, ";".join([self.line_cut_up_key_parameters_name.text(), self.line_cut_up_key_parameters_value.text()]))
        self.item.setText(5, self.text_edit_key_parameters_description.toPlainText())
class CorrectingEditDlg(Ui_Dialog_correcting_edit, QDialog):
    def __init__(self, item, parent=None):
        super().__init__(parent)
        self.setupUi(self)
        self.item=item


        self.line_corrected_parameters_name.setText(item.text(0))
        if item.text(1)==self.radioButton_gradient.text():
            self.radioButton_gradient.setChecked(True)
        elif item.text(1) == self.radioButton_integral.text():
            self.radioButton_integral.setChecked(True)
        else:
            self.radioButton_no_extra_operations.setChecked(True)
        self.text_edit_correcting_expression.setText(item.text(2))

        self.buttonBox.accepted.connect(self.save_changes)

    def save_changes(self):
        self.item.setText(0, self.line_corrected_parameters_name.text())
        if self.radioButton_gradient.isChecked():
            self.item.setText(1, self.radioButton_gradient.text())
        elif self.radioButton_integral.isChecked():
            self.item.setText(1, self.radioButton_integral.text())
        else:
            self.item.setText(1, "")
        self.item.setText(2, self.text_edit_correcting_expression.toPlainText())


class Completers():
    def __init__(self):
        self.reset()

        self.set_help_clusters = set(["min", "max"])
        self.set_help_report = set(["ON_FIRST", "ON_LAST", "OFF_FIRST", "OFF_LAST"])

        self.help_clusters_completer = QCompleter(list(self.set_help_clusters))
        self.help_report_completer = QCompleter(list(self.set_help_report))
        self.help_clusters_report_completer = QCompleter(list(self.set_help_clusters | self.set_help_report))

    def reset(self):
        self.set_clusters = set()
        self.set_report = set()
        self.set_correcting_parameters = set()

        self.clusters_completer = QCompleter(list())
        self.report_completer = QCompleter(list())
        self.clusters_report_completer = QCompleter(list())
        self.graphs_completer = QCompleter(list())

    def update_clusters(self, new_set):
        self.set_clusters = self.set_clusters | set(new_set)
        self.clusters_completer.model().setStringList(list(self.set_clusters | self.set_correcting_parameters))
        self.clusters_report_completer.model().setStringList(list(self.set_clusters | self.set_report | self.set_correcting_parameters))

    def update_report(self, new_set):
        self.set_report = self.set_report | set(new_set)
        self.report_completer.model().setStringList(list(self.set_report))
        self.clusters_report_completer.model().setStringList(list(self.set_clusters | self.set_report))

    def update_correcting_parameters(self, tree_widget):
        number_of_items = tree_widget.topLevelItemCount()
        correcting_parameters=[]
        for i in range(number_of_items):
            item = tree_widget.topLevelItem(i)
            correcting_parameters.append(item.text(0))

        self.set_correcting_parameters = set(correcting_parameters)
        self.clusters_completer.model().setStringList(list(self.set_clusters | self.set_correcting_parameters))
        self.clusters_report_completer.model().setStringList(list(self.set_clusters | self.set_report | self.set_correcting_parameters))

    def update_graphs(self, line):
        list_line = line.text().split(";")
        prefix = ";".join(list_line[:-1])
        previous_set = self.set_clusters | self.set_correcting_parameters
        if prefix=="":
            self.graphs_completer.model().setStringList(list(previous_set))
        else:
            set_with_prefix = [prefix + ";" + name for name in previous_set]
            self.graphs_completer.model().setStringList(set_with_prefix)


class Worker(QRunnable):
    def __init__(self, fn, *args, **kwargs):
        super(Worker, self).__init__()
        self.fn = fn
        self.args = args
        self.kwargs = kwargs
        self.signals = WorkerSignals()
    @Slot()
    def run(self):
        try:
            result = self.fn(*self.args, **self.kwargs)
        except:
            traceback.print_exc()
            exctype, value = sys.exc_info()[:2]
            self.signals.error.emit((exctype, value, traceback.format_exc()))
        else:
            self.signals.result.emit(result)
        finally:
            self.signals.finished.emit()

class WorkerSignals(QObject):
    finished = Signal()
    error = Signal(tuple)
    result = Signal(object)


if __name__ == '__main__':
    app = QApplication()
    window = MainWindow()
    window.show()
    sys.exit(app.exec())


